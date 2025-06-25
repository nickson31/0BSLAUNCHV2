# -*- coding: utf-8 -*-
"""
0Bullshit Backend v2.0 - CORS ARREGLADO DEFINITIVAMENTE
Sistema de 60 bots, memoria neuronal, créditos y chat sessions
VERSIÓN FINAL SIN PROBLEMAS DE CORS
"""

# ==============================================================================
#           IMPORTS
# ==============================================================================

print("1. Loading libraries...")
from flask import Flask, request, jsonify, session, redirect, url_for, make_response
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import pandas as pd
import google.generativeai as genai
import json
import ast
import re
import warnings
import random
import time
import os
import sqlalchemy
from datetime import datetime, timedelta
import uuid
from sqlalchemy import text
import secrets
import jwt
from functools import wraps
import hashlib
import bcrypt
import requests
import tempfile
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage
import PyPDF2
import docx
import mimetypes
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# Google Auth
from google.oauth2 import id_token
from google.auth.transport import requests as google_requests

# ==============================================================================
#           CONFIGURATION
# ==============================================================================

print("2. Configuring application...")
app = Flask(__name__)

# ==================== CORS KILLER - CONFIGURACIÓN NUCLEAR ====================
print("🛡️ Configurando CORS KILLER...")

CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": "*",
        "allow_headers": "*"
    }
})

# Rate Limiter Configuration
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)

app.secret_key = os.environ.get('JWT_SECRET', secrets.token_hex(16))
warnings.filterwarnings('ignore')

# Environment Variables
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
DATABASE_URL = os.environ.get("DATABASE_URL")
JWT_SECRET = os.environ.get("JWT_SECRET", secrets.token_hex(32))
UNIPILE_API_KEY = os.environ.get("UNIPILE_API_KEY")
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET")
FRONTEND_URL = os.environ.get("FRONTEND_URL", "https://v0-0-bull-shit.vercel.app")

# Debug info
print(f"🔐 GOOGLE_CLIENT_ID configured: {'✅' if GOOGLE_CLIENT_ID else '❌'}")
print(f"🗄️ DATABASE_URL configured: {'✅' if DATABASE_URL else '❌'}")
print(f"🔑 JWT_SECRET configured: {'✅' if JWT_SECRET else '❌'}")
print(f"🤖 GEMINI_API_KEY configured: {'✅' if GEMINI_API_KEY else '❌'}")

# Verificar configuración crítica
if not GEMINI_API_KEY:
    print("❌ FATAL: GEMINI_API_KEY not found.")
if not DATABASE_URL:
    print("❌ FATAL: DATABASE_URL not found.")
if not GOOGLE_CLIENT_ID:
    print("⚠️ WARNING: GOOGLE_CLIENT_ID not found - Google Auth will not work")

# Configure AI APIs
try:
    genai.configure(api_key=GEMINI_API_KEY)
    MODEL_NAME = "gemini-2.0-flash"
    print("✅ Gemini API configured.")
except Exception as e:
    print(f"❌ ERROR configuring Gemini: {e}")

# Connect to Database
try:
    engine = sqlalchemy.create_engine(DATABASE_URL)
    print("✅ Database connection established.")
except Exception as e:
    print(f"❌ ERROR connecting to database: {e}")
    engine = None

# ==============================================================================
#           CORS MIDDLEWARE - CONFIGURACIÓN DEFINITIVA
# ==============================================================================

@app.before_request
def handle_options():
    if request.method == "OPTIONS":
        response = make_response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "*" 
        response.headers["Access-Control-Allow-Methods"] = "*"
        return response

@app.after_request
def add_cors_headers(response):
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Headers"] = "*"
    response.headers["Access-Control-Allow-Methods"] = "*"
    return response

# ==============================================================================
#           CONSTANTS AND CONFIGURATIONS
# ==============================================================================

# Credit costs por acción
CREDIT_COSTS = {
    # Bots básicos (todos los planes)
    "basic_bot": 5,
    "advanced_bot": 15,
    "expert_bot": 25,
    "document_generation": 50,
    
    # Growth plan only
    "investor_search_result": 10,
    "employee_search_result": 8,
    
    # Pro plan only
    "template_generation": 20,
    "unipile_message": 5,
    "automated_sequence": 50,
    
    # Premium features
    "market_analysis": 100,
    "business_model": 150,
    "pitch_deck": 200
}

# Planes de suscripción
SUBSCRIPTION_PLANS = {
    "free": {
        "name": "Free",
        "credits_monthly": 10000,
        "launch_credits": 10000,
        "features": {
            "bots_access": True,
            "investor_search": False,
            "employee_search": False,
            "outreach_templates": False,
            "unlimited_docs": False,
            "neural_memory": False
        }
    },
    "growth": {
        "name": "Growth",
        "credits_monthly": 100000,
        "launch_credits": 1000000,
        "features": {
            "bots_access": True,
            "investor_search": True,
            "employee_search": True,
            "outreach_templates": False,
            "unlimited_docs": True,
            "neural_memory": True
        }
    },
    "pro": {
        "name": "Pro Outreach",
        "credits_monthly": 50000,
        "launch_credits": 1000000,
        "features": {
            "bots_access": True,
            "investor_search": True,
            "employee_search": True,
            "outreach_templates": True,
            "unlimited_docs": True,
            "neural_memory": True,
            "unipile_integration": True
        }
    }
}

# ==============================================================================
#           AUTHENTICATION & USER MANAGEMENT
# ==============================================================================

def get_bot_credit_cost(bot_id):
    """Obtiene el costo en créditos de un bot"""
    return CREDIT_COSTS.get(bot_id, 5)

def hash_password(password):
    """Hash password usando bcrypt"""
    if not password:
        return None
    salt = bcrypt.gensalt()
    return bcrypt.hashpw(password.encode('utf-8'), salt).decode('utf-8')

def verify_password(password, hashed):
    """Verifica password contra hash"""
    if not password or not hashed:
        return False
    try:
        return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))
    except Exception as e:
        print(f"Error verifying password: {e}")
        return False

def generate_jwt_token(user_id):
    """Genera JWT token para el usuario"""
    try:
        payload = {
            'user_id': str(user_id),
            'exp': datetime.utcnow() + timedelta(days=7),
            'iat': datetime.utcnow()
        }
        return jwt.encode(payload, JWT_SECRET, algorithm='HS256')
    except Exception as e:
        print(f"Error generating JWT: {e}")
        return None

def verify_jwt_token(token):
    """Verifica y decodifica JWT token"""
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
        return payload['user_id']
    except jwt.ExpiredSignatureError:
        print("Token expired")
        return None
    except jwt.InvalidTokenError:
        print("Invalid token")
        return None
    except Exception as e:
        print(f"Error verifying token: {e}")
        return None

def require_plan(required_plan):
    """
    Decorator para verificar que el usuario tiene el plan requerido
    """
    def decorator(f):
        @wraps(f)
        def decorated_function(user, *args, **kwargs):
            user_plan = user.get('plan', 'free')
            
            # Jerarquía de planes: free < growth < pro
            plan_hierarchy = {'free': 1, 'growth': 2, 'pro': 3}
            
            user_level = plan_hierarchy.get(user_plan, 1)
            required_level = plan_hierarchy.get(required_plan, 1)
            
            if user_level < required_level:
                return jsonify({
                    'error': 'plan_upgrade_required',
                    'current_plan': user_plan,
                    'required_plan': required_plan,
                    'message': f'Este feature requiere plan {required_plan}'
                }), 403
            
            return f(user, *args, **kwargs)
        return decorated_function
    return decorator

def ml_investor_search(query, user_preferences, max_results=20):
    """
    Función principal de búsqueda ML de inversores - VERSIÓN CORREGIDA
    """
    try:
        print(f"🔍 Iniciando búsqueda ML: '{query}'")
        
        # Verificar que engine esté disponible
        if not engine:
            print("❌ Engine de base de datos no disponible")
            return {
                "error": "Database connection not available",
                "success": False,
                "query": query
            }
        
        # Crear instancia del motor de búsqueda
        search_engine = InvestorSearchSimple(engine)
        
        # Ejecutar búsqueda
        resultado = search_engine.buscar_inversores(query)
        
        print(f"✅ Búsqueda completada: {len(resultado.get('results', []))} resultados")
        return resultado
        
    except Exception as e:
        print(f"❌ Error en ml_investor_search: {e}")
        import traceback
        traceback.print_exc()
        return {
            "error": f"Search failed: {str(e)}",
            "success": False,
            "query": query,
            "details": str(e)
        }
        
def require_auth(f):
    """Decorator para endpoints que requieren autenticación"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = request.headers.get('Authorization')
        if not token:
            return jsonify({'error': 'No authorization token provided'}), 401
        
        if token.startswith('Bearer '):
            token = token[7:]
        
        user_id = verify_jwt_token(token)
        if not user_id:
            return jsonify({'error': 'Invalid or expired token'}), 401
        
        # Obtener usuario de la base de datos
        user = get_user_by_id(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        return f(user, *args, **kwargs)
    return decorated_function

def get_user_by_id(user_id):
    """Obtiene usuario por ID - CORREGIDO para tu esquema EXACTO"""
    try:
        with engine.connect() as conn:
            result = conn.execute(
                text("""
                    SELECT id, email, password_hash, first_name, last_name, 
                           plan, credits, auth_provider, google_id, 
                           is_active, is_admin, created_at, updated_at
                    FROM users 
                    WHERE id = :user_id
                """),
                {"user_id": user_id}
            ).fetchone()
            
            if result:
                return {
                    'id': str(result[0]),
                    'email': result[1],
                    'password_hash': result[2],
                    'first_name': result[3] or '',
                    'last_name': result[4] or '',
                    'subscription_plan': result[5] or 'free',  # Mapear "plan" a "subscription_plan"
                    'plan': result[5] or 'free',  # También mantener plan original
                    'credits': result[6] or 100,
                    'auth_provider': result[7] or 'manual',
                    'google_id': result[8],
                    'is_active': result[9] if result[9] is not None else True,
                    'is_admin': result[10] if result[10] is not None else False,
                    'created_at': result[11],
                    'updated_at': result[12]
                }
            return None
    except Exception as e:
        print(f"❌ Error getting user by ID: {e}")
        return None

def get_user_by_email(email):
    """Obtiene usuario por email - CORREGIDO para tu esquema EXACTO"""
    try:
        with engine.connect() as conn:
            result = conn.execute(
                text("""
                    SELECT id, email, password_hash, first_name, last_name, 
                           plan, credits, auth_provider, google_id, 
                           is_active, is_admin, created_at, updated_at
                    FROM users 
                    WHERE email = :email
                """),
                {"email": email}
            ).fetchone()
            
            if result:
                return {
                    'id': str(result[0]),
                    'email': result[1],
                    'password_hash': result[2],
                    'first_name': result[3] or '',
                    'last_name': result[4] or '',
                    'subscription_plan': result[5] or 'free',  # Mapear "plan" a "subscription_plan"
                    'plan': result[5] or 'free',  # También mantener plan original
                    'credits': result[6] or 100,
                    'auth_provider': result[7] or 'manual',
                    'google_id': result[8],
                    'is_active': result[9] if result[9] is not None else True,
                    'is_admin': result[10] if result[10] is not None else False,
                    'created_at': result[11],
                    'updated_at': result[12]
                }
            return None
    except Exception as e:
        print(f"❌ Error getting user by email: {e}")
        return None

def get_user_by_google_id(google_id):
    """Obtiene usuario por Google ID"""
    try:
        with engine.connect() as conn:
            result = conn.execute(
                text("SELECT * FROM users WHERE google_id = :google_id"),
                {"google_id": google_id}
            ).fetchone()
            
            if result:
                return {
                    'id': str(result[0]),
                    'email': result[1],
                    'password_hash': result[2],
                    'first_name': result[3] or '',
                    'last_name': result[4] or '',
                    'subscription_plan': result[5] or 'free',
                    'credits': result[6] or 100,
                    'auth_provider': result[7] or 'manual',
                    'google_id': result[8],
                    'is_active': result[9] if len(result) > 9 else True,
                    'is_admin': result[10] if len(result) > 10 else False,
                    'created_at': result[11] if len(result) > 11 else None
                }
            return None
    except Exception as e:
        print(f"Error getting user by Google ID: {e}")
        return None

def create_user(email, password=None, first_name="", last_name="", auth_provider="manual", google_id=None):
    """Crea nuevo usuario"""
    try:
        user_id = str(uuid.uuid4())
        password_hash = hash_password(password) if password else None
        
        # Creditos iniciales por plan
        initial_credits = SUBSCRIPTION_PLANS['free']['launch_credits']
        
        with engine.connect() as conn:
            conn.execute(
                text("""
                    INSERT INTO users (
                        id, email, password_hash, first_name, last_name,
                        plan, credits, auth_provider, google_id, 
                        is_active, is_admin, created_at, updated_at
                    ) VALUES (
                        :id, :email, :password_hash, :first_name, :last_name,
                        'free', :credits, :auth_provider, :google_id, 
                        true, false, NOW(), NOW()
                    )
                """),
                {
                    "id": user_id,
                    "email": email,
                    "password_hash": password_hash,
                    "first_name": first_name,
                    "last_name": last_name,
                    "credits": initial_credits,
                    "auth_provider": auth_provider,
                    "google_id": google_id
                }
            )
            conn.commit()
            print(f"✅ User created successfully: {user_id}")
            return user_id
    except Exception as e:
        print(f"❌ Error creating user: {e}")
        return None

def validate_password_strength(password):
    """Valida que la contraseña cumpla con los requisitos"""
    if not password:
        return False, "Password is required"
    
    if len(password) < 8:
        return False, "Password must be at least 8 characters"
    
    if not any(c.isupper() for c in password):
        return False, "Password must contain at least one uppercase letter"
    
    if not any(c.islower() for c in password):
        return False, "Password must contain at least one lowercase letter"
    
    if not any(c.isdigit() for c in password):
        return False, "Password must contain at least one number"
    
    return True, "Password is valid"

def get_user_credits(user_id):
    """Obtiene créditos del usuario"""
    try:
        with engine.connect() as conn:
            result = conn.execute(
                text("SELECT credits FROM users WHERE id = :user_id"),
                {"user_id": user_id}
            ).fetchone()
            return result[0] if result else 0
    except Exception as e:
        print(f"Error getting credits: {e}")
        return 0

def charge_credits(user_id, amount):
    """Cobra créditos al usuario - PERFECTO para tu esquema"""
    try:
        print(f"💸 Cobrando {amount} créditos al usuario {user_id}")
        
        with engine.connect() as conn:
            # Verificar créditos actuales
            current_result = conn.execute(
                text("SELECT credits FROM users WHERE id = :user_id"),
                {"user_id": user_id}
            ).fetchone()
            
            if not current_result:
                print(f"❌ Usuario {user_id} no encontrado")
                return None
            
            current_credits = current_result[0] or 0
            print(f"💰 Créditos actuales: {current_credits}")
            
            if current_credits < amount:
                print(f"❌ Créditos insuficientes: {current_credits} < {amount}")
                return None
            
            # Hacer el cobro - updated_at existe en tu tabla
            result = conn.execute(
                text("""
                    UPDATE users 
                    SET credits = credits - :amount,
                        updated_at = NOW()
                    WHERE id = :user_id
                    RETURNING credits
                """),
                {"user_id": user_id, "amount": amount}
            ).fetchone()
            
            if result:
                new_credits = result[0]
                conn.commit()
                
                print(f"✅ Cobro exitoso: {current_credits} -> {new_credits}")
                
                # Log transaction - tu tabla credit_transactions existe
                try:
                    log_credit_transaction(user_id, -amount, 'charge', 'Bot usage')
                except Exception as log_error:
                    print(f"⚠️ Error logging transaction: {log_error}")
                
                return new_credits
            else:
                print(f"❌ Error en UPDATE de créditos")
                return None
                
    except Exception as e:
        print(f"❌ Error charging credits: {e}")
        import traceback
        traceback.print_exc()
        return None
        
def add_credits(user_id, amount, reason='purchase'):
    """Agrega créditos al usuario"""
    try:
        with engine.connect() as conn:
            result = conn.execute(
                text("""
                    UPDATE users 
                    SET credits = credits + :amount 
                    WHERE id = :user_id
                    RETURNING credits
                """),
                {"user_id": user_id, "amount": amount}
            ).fetchone()
            
            if result:
                log_credit_transaction(user_id, amount, 'credit', reason)
                conn.commit()
                return result[0]
            return None
    except Exception as e:
        print(f"Error adding credits: {e}")
        return None

def log_credit_transaction(user_id, amount, transaction_type, description):
    """Registra transacción de créditos - PERFECTO para tu esquema"""
    try:
        with engine.connect() as conn:
            conn.execute(
                text("""
                    INSERT INTO credit_transactions (
                        id, user_id, amount, transaction_type, description, created_at
                    ) VALUES (
                        :id, :user_id, :amount, :type, :description, NOW()
                    )
                """),
                {
                    "id": str(uuid.uuid4()),
                    "user_id": user_id,
                    "amount": amount,
                    "type": transaction_type,
                    "description": description
                }
            )
            conn.commit()
    except Exception as e:
        print(f"Error logging transaction: {e}")

def has_sufficient_credits(user_id, required_amount):
    """Verifica si el usuario tiene suficientes créditos"""
    try:
        current_credits = get_user_credits(user_id)
        return current_credits >= required_amount
    except Exception as e:
        print(f"Error checking credits: {e}")
        return False

def update_subscription_plan(user_id, new_plan):
    """Actualiza el plan de suscripción del usuario"""
    try:
        if new_plan not in SUBSCRIPTION_PLANS:
            return False
            
        with engine.connect() as conn:
            # Update user's subscription plan
            conn.execute(
                text("""
                    UPDATE users 
                    SET plan = :plan,
                        credits = credits + :additional_credits,
                        updated_at = NOW()
                    WHERE id = :user_id
                """),
                {
                    "user_id": user_id,
                    "plan": new_plan,
                    "additional_credits": SUBSCRIPTION_PLANS[new_plan]['launch_credits']
                }
            )
            conn.commit()
            return True
    except Exception as e:
        print(f"Error updating subscription: {e}")
        return False

def verify_google_access_token(access_token):
    """Verifica Google access token"""
    try:
        print(f"🔍 Verifying Google access token...")
        
        response = requests.get(
            f'https://www.googleapis.com/oauth2/v2/userinfo?access_token={access_token}',
            timeout=10
        )
        
        print(f"📡 Google API response status: {response.status_code}")
        
        if response.status_code != 200:
            print(f"❌ Google API error: {response.text}")
            return None
        
        user_info = response.json()
        print(f"👤 User info received: {user_info}")
        
        return {
            'google_id': user_info.get('id'),
            'email': user_info.get('email'),
            'first_name': user_info.get('given_name', ''),
            'last_name': user_info.get('family_name', ''),
            'picture': user_info.get('picture', ''),
            'verified_email': user_info.get('verified_email', False)
        }
        
    except Exception as e:
        print(f"❌ Error verifying Google token: {e}")
        return None

# ==============================================================================
#           NEURAL MEMORY SYSTEM
# ==============================================================================

def init_neural_memory(user_id):
    """Inicializa memoria neuronal para usuario"""
    try:
        with engine.connect() as conn:
            # Verificar si ya existe
            existing = conn.execute(
                text("SELECT id FROM neural_memory WHERE user_id = :user_id"),
                {"user_id": user_id}
            ).fetchone()
            
            if not existing:
                conn.execute(
                    text("""
                        INSERT INTO neural_memory (user_id, memory_data, created_at, updated_at)
                        VALUES (:user_id, '{}', NOW(), NOW())
                    """),
                    {"user_id": user_id}
                )
                conn.commit()
                print(f"✅ Neural memory initialized for user: {user_id}")
    except Exception as e:
        print(f"❌ Error initializing neural memory: {e}")

def save_neural_interaction(user_id, interaction_data):
    """
    Guarda interacción en memoria neuronal - OPTIMIZADO para tu esquema
    Todas las columnas existen en tu DB, así que esta función es PERFECTA
    """
    try:
        interaction_id = str(uuid.uuid4())
        
        with engine.connect() as conn:
            conn.execute(
                text("""
                    INSERT INTO neural_interactions (
                        id, user_id, bot_used, user_input, bot_output, 
                        credits_charged, context_data, session_id, project_id, created_at
                    ) VALUES (
                        :id, :user_id, :bot_used, :user_input, :bot_output,
                        :credits_charged, :context_data, :session_id, :project_id, NOW()
                    )
                """),
                {
                    "id": interaction_id,
                    "user_id": user_id,
                    "bot_used": interaction_data.get('bot', 'unknown'),
                    "user_input": interaction_data.get('input', ''),
                    "bot_output": interaction_data.get('response', ''),
                    "credits_charged": interaction_data.get('credits_used', 0),
                    "context_data": json.dumps(interaction_data.get('context', {})),
                    "session_id": interaction_data.get('session_id'),
                    "project_id": interaction_data.get('project_id')
                }
            )
            conn.commit()
            
            # Si es la primera interacción de la sesión, generar título
            session_id = interaction_data.get('session_id')
            if session_id:
                check_and_generate_session_title(session_id, user_id, interaction_data)
            
            print(f"✅ Interaction saved: {interaction_id}")
    
    except Exception as e:
        print(f"❌ Error saving interaction: {e}")
        import traceback
        traceback.print_exc()

def check_and_generate_session_title(session_id, user_id, interaction_data):
    """Genera título para la sesión si es la primera interacción"""
    try:
        with engine.connect() as conn:
            # Verificar si ya existe título para esta sesión
            existing_title = conn.execute(
                text("""
                    SELECT session_title FROM neural_interactions 
                    WHERE session_id = :session_id AND user_id = :user_id 
                    AND session_title IS NOT NULL
                    LIMIT 1
                """),
                {"session_id": session_id, "user_id": user_id}
            ).fetchone()
            
            if existing_title:
                return  # Ya tiene título
            
            # Generar título con Gemini
            user_input = interaction_data.get('input', '')
            bot_response = interaction_data.get('response', '')
            
            title = generate_chat_title_with_gemini(user_input, bot_response)
            
            # Actualizar la interacción con el título
            conn.execute(
                text("""
                    UPDATE neural_interactions 
                    SET session_title = :title 
                    WHERE session_id = :session_id AND user_id = :user_id
                """),
                {"title": title, "session_id": session_id, "user_id": user_id}
            )
            conn.commit()
            
            print(f"✅ Session title generated: {title}")
            
    except Exception as e:
        print(f"❌ Error generating session title: {e}")

def generate_chat_title_with_gemini(user_input, bot_response):
    """Genera título inteligente usando Gemini - MULTIIDIOMA"""
    try:
        # Primero detectar el idioma
        user_language = detect_user_language_with_gemini(user_input)
        
        # Prompts en diferentes idiomas
        title_prompts = {
            'es': f"""
            Genera un título corto y descriptivo (máximo 5 palabras) para esta conversación.
            El título debe capturar la esencia de lo que el usuario está preguntando o trabajando.
            
            Usuario preguntó: {user_input[:200]}
            Asistente respondió sobre: {bot_response[:200]}
            
            Ejemplos de buenos títulos:
            - "Pitch Deck para FinTech"
            - "Buscar Inversores Seed"
            - "Plan Marketing SaaS"
            - "Análisis Competencia EdTech"
            - "Modelo Financiero B2B"
            
            Responde SOLO con el título en el idioma de input/prompt del usuario, si es español, español, si es chino, chino, si es inglés, inglés, sin comillas ni explicaciones:
            """,
            
            'en': f"""
            Generate a short and descriptive title (maximum 5 words) for this conversation.
            The title should capture the essence of what the user is asking or working on.
            
            User asked: {user_input[:200]}
            Assistant responded about: {bot_response[:200]}
            
            Examples of good titles:
            - "FinTech Pitch Deck"
            - "Find Seed Investors"
            - "SaaS Marketing Plan"
            - "EdTech Competition Analysis"
            - "B2B Financial Model"
            
            Respond ONLY with the title in the language of the user's input/prompt, if the conversation is in English, English, if the user talks in Spanish, Spanish, no quotes or explanations:
            """
        }
        
        # Use appropriate prompt based on language
        title_prompt = title_prompts.get(user_language, title_prompts['en'])
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            title_prompt,
            generation_config={
                "temperature": 0.3,
                "max_output_tokens": 20,
            }
        )
        
        title = response.text.strip()
        
        # Limpiar el título
        title = title.replace('"', '').replace("'", '').strip()
        
        # Validar longitud
        if len(title) > 50:
            title = title[:47] + "..."
        
        # Títulos por defecto según idioma
        default_titles = {
            'es': "Nueva Conversación",
            'en': "New Conversation",
            'fr': "Nouvelle Conversation",
            'de': "Neues Gespräch",
            'pt': "Nova Conversa",
            'it': "Nuova Conversazione"
        }
        
        return title if title else default_titles.get(user_language, "New Conversation")
        
    except Exception as e:
        print(f"❌ Error generating title: {e}")
        return "Nueva Conversación"

def generate_simple_title_from_message(first_message, last_message):
    """Genera título simple cuando no hay uno guardado - FALLBACK"""
    try:
        # Usar primeras palabras del primer mensaje o último
        if first_message and len(first_message) > 10:
            # Tomar primeras 5-7 palabras
            words = first_message.split()[:6]
            title = " ".join(words)
            if len(first_message) > len(title):
                title += "..."
        elif last_message and len(last_message) > 10:
            words = last_message.split()[:6]
            title = " ".join(words)
            if len(last_message) > len(title):
                title += "..."
        else:
            title = "Nueva Conversación"
        
        # Limpiar caracteres especiales
        title = title.replace('\n', ' ').replace('\r', ' ')
        
        # Limitar longitud
        if len(title) > 50:
            title = title[:47] + "..."
            
        return title
    except Exception as e:
        print(f"❌ Error generating simple title: {e}")
        return "Nueva Conversación"

def get_neural_memory(user_id):
    """Obtiene memoria neuronal del usuario"""
    try:
        with engine.connect() as conn:
            result = conn.execute(
                text("""
                    SELECT memory_data FROM neural_memory 
                    WHERE user_id = :user_id
                """),
                {"user_id": user_id}
            ).fetchone()
            
            if result and result[0]:
                return json.loads(result[0])
            return {}
    except Exception as e:
        print(f"Error getting neural memory: {e}")
        return {}

def extract_and_update_project_memory(user_id, project_id, user_input, bot_response):
    """
    Versión corregida que maneja JSON correctamente Y GUARDA INTENCIÓN DETECTADA
    """
    try:
        # Obtener memoria actual del proyecto
        with engine.connect() as conn:
            result = conn.execute(
                text("""
                    SELECT kpi_data FROM projects 
                    WHERE id = :project_id AND user_id = :user_id
                """),
                {"project_id": project_id, "user_id": user_id}
            ).fetchone()
            
            if not result:
                return False
            
            # Manejar JSON correctamente
            try:
                if result[0] is None:
                    current_memory = {}
                elif isinstance(result[0], str):
                    current_memory = json.loads(result[0])
                elif isinstance(result[0], dict):
                    current_memory = result[0]
                else:
                    current_memory = {}
            except (json.JSONDecodeError, TypeError):
                current_memory = {}
        
        # Actualizar memoria con más contexto
        updated_memory = current_memory.copy()
        updated_memory["last_updated"] = datetime.now().isoformat()
        
        # NUEVO: Guardar última intención detectada y contexto de búsqueda
        if "inversores" in bot_response or "investors" in bot_response:
            updated_memory["last_investor_search"] = {
                "timestamp": datetime.now().isoformat(),
                "query": user_input,
                "detected_intent": "investor_search",
                "response_preview": bot_response[:200]
            }
        
        # Detectar si es un Startup Profile
        if "startup profile" in user_input.lower() or "**sector:**" in user_input.lower():
            updated_memory["startup_profile_provided"] = {
                "timestamp": datetime.now().isoformat(),
                "profile_text": user_input,
                "parsed_data": extract_business_info(user_input)
            }
        
        # Guardar memoria actualizada - CONVERTIR A STRING JSON
        with engine.connect() as conn:
            conn.execute(
                text("""
                    UPDATE projects 
                    SET kpi_data = :memory_data, updated_at = NOW()
                    WHERE id = :project_id AND user_id = :user_id
                """),
                {
                    "memory_data": json.dumps(updated_memory),  # CONVERTIR A STRING
                    "project_id": project_id,
                    "user_id": user_id
                }
            )
            conn.commit()
        
        print(f"✅ Project memory updated for project {project_id}")
        return True
        
    except Exception as e:
        print(f"❌ Error updating project memory: {e}")
        return False
        
def extract_business_info(text):
    """
    Extrae información de negocio del texto usando Gemini - MULTIIDIOMA
    """
    try:
        # Detectar idioma primero
        detected_language = detect_user_language_with_gemini(text[:200])
        
        extraction_prompts = {
            'es': f"""
            Analiza el siguiente texto y extrae SOLO la información de negocio específica que mencione el usuario.
            NO inventes información que no esté explícitamente mencionada.
            
            Texto: {text}
            
            Devuelve SOLO un JSON con esta estructura (solo incluye campos que tengan información real):
            {{
                "startup_name": "nombre si se menciona",
                "industry": "industria específica si se menciona", 
                "stage": "etapa si se menciona (idea, mvp, seed, series_a, etc)",
                "business_model": "modelo de negocio si se describe",
                "target_market": "mercado objetivo si se especifica",
                "problem_solving": "problema que resuelve si se explica",
                "revenue_model": "como genera dinero si se menciona",
                "team_size": "tamaño del equipo si se menciona",
                "location": "ubicación si se especifica",
                "funding_raised": "dinero levantado si se menciona",
                "funding_needed": "dinero que necesita si se menciona",
                "competitors": ["competidores si se mencionan"],
                "key_metrics": "métricas clave si se mencionan",
                "current_challenges": ["retos actuales si se mencionan"],
                "business_type": "real_startup, side_project, o idea_stage"
            }}
            
            Si no hay información específica, devuelve {{}}.
            """,
            
            'en': f"""
            Analyze the following text and extract ONLY the specific business information mentioned by the user.
            DO NOT invent information that is not explicitly mentioned.
            
            Text: {text}
            
            Return ONLY a JSON with this structure (only include fields that have real information):
            {{
                "startup_name": "name if mentioned",
                "industry": "specific industry if mentioned", 
                "stage": "stage if mentioned (idea, mvp, seed, series_a, etc)",
                "business_model": "business model if described",
                "target_market": "target market if specified",
                "problem_solving": "problem being solved if explained",
                "revenue_model": "how it makes money if mentioned",
                "team_size": "team size if mentioned",
                "location": "location if specified",
                "funding_raised": "money raised if mentioned",
                "funding_needed": "money needed if mentioned",
                "competitors": ["competitors if mentioned"],
                "key_metrics": "key metrics if mentioned",
                "current_challenges": ["current challenges if mentioned"],
                "business_type": "real_startup, side_project, or idea_stage"
            }}
            
            If there's no specific information, return {{}}.
            """
        }
        
        extraction_prompt = extraction_prompts.get(detected_language, extraction_prompts['en'])
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            extraction_prompt,
            generation_config={
                "temperature": 0.1,
                "max_output_tokens": 1000,
            }
        )
        
        # Extraer JSON de la respuesta
        response_text = response.text.strip()
        if response_text.startswith('```json'):
            response_text = response_text[7:-3]
        elif response_text.startswith('```'):
            response_text = response_text[3:-3]
        
        try:
            extracted_data = json.loads(response_text)
            return extracted_data
        except json.JSONDecodeError:
            print(f"⚠️ Could not parse JSON from Gemini: {response_text}")
            return {}
        
    except Exception as e:
        print(f"❌ Error extracting business info: {e}")
        return {}

def merge_memory_data(current_memory, new_info):
    """
    Combina memoria actual con nueva información inteligentemente
    """
    try:
        # Inicializar estructura si no existe
        if not current_memory:
            current_memory = {
                "business_context": {},
                "chat_history_summary": [],
                "investor_preferences": {},
                "document_history": [],
                "last_updated": datetime.now().isoformat()
            }
        
        # Actualizar contexto de negocio
        business_context = current_memory.get("business_context", {})
        
        for key, value in new_info.items():
            if value and value != "":  # Solo actualizar si hay valor real
                if key == "competitors" and isinstance(value, list):
                    # Combinar listas de competidores
                    existing = business_context.get(key, [])
                    business_context[key] = list(set(existing + value))
                elif key == "current_challenges" and isinstance(value, list):
                    # Combinar listas de retos
                    existing = business_context.get(key, [])
                    business_context[key] = list(set(existing + value))
                else:
                    # Actualizar valor simple
                    business_context[key] = value
        
        current_memory["business_context"] = business_context
        current_memory["last_updated"] = datetime.now().isoformat()
        
        return current_memory
        
    except Exception as e:
        print(f"❌ Error merging memory data: {e}")
        return current_memory or {}

def detect_user_language_with_gemini(text):
    """Detecta el idioma del usuario usando Gemini - VERSIÓN INTELIGENTE"""
    try:
        if not text or len(text.strip()) < 3:
            return 'en'  # Default to English for very short inputs
        
        detection_prompt = f"""
        Detect the language of this text and respond with ONLY the ISO 639-1 language code (2 letters).
        
        Text: "{text}"
        
        Examples of responses:
        - For English text: en
        - For Spanish text: es
        - For French text: fr
        - For German text: de
        - For Portuguese text: pt
        - For Italian text: it
        
        Respond with ONLY the 2-letter code, nothing else:
        """
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            detection_prompt,
            generation_config={
                "temperature": 0.1,  # Very low temperature for consistent detection
                "max_output_tokens": 10,
            }
        )
        
        detected_language = response.text.strip().lower()[:2]  # Get only first 2 letters
        
        # Validate the response
        valid_languages = ['en', 'es', 'fr', 'de', 'pt', 'it', 'nl', 'ru', 'zh', 'ja', 'ko', 'ar']
        
        if detected_language in valid_languages:
            print(f"🌍 Detected language: {detected_language} for text: '{text[:50]}...'")
            return detected_language
        else:
            print(f"⚠️ Invalid language code detected: {detected_language}, defaulting to English")
            return 'en'
            
    except Exception as e:
        print(f"❌ Error detecting language with Gemini: {e}")
        # Fallback to simple detection if Gemini fails
        return detect_user_language_simple_fallback(text)

def detect_user_language_simple_fallback(text):
    """Detección simple de idioma como fallback"""
    spanish_keywords = ['hola', 'como', 'que', 'para', 'por', 'con', 'esto', 'esta', 'quiero', 'necesito', 'puedo', 'ayuda']
    english_keywords = ['hello', 'how', 'what', 'for', 'with', 'this', 'want', 'need', 'can', 'help', 'please']
    
    text_lower = text.lower()
    
    spanish_count = sum(1 for word in spanish_keywords if word in text_lower)
    english_count = sum(1 for word in english_keywords if word in text_lower)
    
    if spanish_count > english_count:
        return 'es'
    else:
        return 'en'

def detect_user_language(text):
    """Función principal de detección de idioma - usa Gemini con fallback"""
    return detect_user_language_with_gemini(text)
        
def get_enhanced_context_for_chat(user, session_id, project_id, data):
    """Obtiene contexto mejorado para el chat - CON DETECCIÓN DE IDIOMA GEMINI"""
    try:
        # Obtener contexto del proyecto
        project_context = get_project_context_for_chat(user['id'], project_id)
        
        # DETECTAR IDIOMA CON GEMINI
        user_message = data.get('message', '')
        detected_language = detect_user_language_with_gemini(user_message)
        
        # Contexto base
        enhanced_context = {
            'user_id': user['id'],
            'user_plan': user.get('plan', 'free'),
            'session_id': session_id,
            'project_id': project_id,
            'user_credits_before': get_user_credits(user['id']),
            'user_language': detected_language,  # IDIOMA DETECTADO POR GEMINI
            'user_message': user_message,
            **data.get('context', {}),
            **project_context
        }
        
        return enhanced_context
    except Exception as e:
        print(f"❌ Error getting enhanced context: {e}")
        return {
            'user_id': user['id'],
            'user_plan': user.get('plan', 'free'),
            'session_id': session_id,
            'project_id': project_id,
            'user_language': 'en'  # Default language
        }
        
def get_project_context_for_chat(user_id, project_id):
    """
    Obtiene el contexto completo del proyecto para el chat - OPTIMIZADO
    """
    try:
        with engine.connect() as conn:
            # Obtener información del proyecto
            project_result = conn.execute(
                text("""
                    SELECT project_name, project_description, kpi_data, status, created_at
                    FROM projects 
                    WHERE id = :project_id AND user_id = :user_id
                """),
                {"project_id": project_id, "user_id": user_id}
            ).fetchone()
            
            if not project_result:
                return {}
            
            # Obtener últimas interacciones del proyecto
            recent_chats = conn.execute(
                text("""
                    SELECT user_input, bot_output, created_at 
                    FROM neural_interactions 
                    WHERE user_id = :user_id AND project_id = :project_id 
                    ORDER BY created_at DESC 
                    LIMIT 10
                """),
                {"user_id": user_id, "project_id": project_id}
            ).fetchall()
            
            # Obtener documentos del proyecto
            project_docs = conn.execute(
                text("""
                    SELECT document_type, title, created_at 
                    FROM generated_documents 
                    WHERE user_id = :user_id 
                    AND metadata::jsonb->>'project_id' = :project_id
                    ORDER BY created_at DESC
                """),
                {"user_id": user_id, "project_id": project_id}
            ).fetchall()
        
        # Procesar memoria del proyecto
        project_memory = json.loads(project_result[2]) if project_result[2] else {}
        business_context = project_memory.get("business_context", {})
        
        # Crear resumen de chats recientes
        recent_context = []
        for chat in recent_chats[:5]:  # Últimos 5 chats
            recent_context.append({
                "user_said": chat[0][:100] + "..." if len(chat[0]) > 100 else chat[0],
                "assistant_responded": chat[1][:100] + "..." if len(chat[1]) > 100 else chat[1],
                "when": chat[2].strftime("%Y-%m-%d")
            })
        
        # Crear lista de documentos
        documents_created = [
            {
                "type": doc[0],
                "title": doc[1],
                "created": doc[2].strftime("%Y-%m-%d")
            }
            for doc in project_docs
        ]
        
        return {
            "project_name": project_result[0],
            "project_description": project_result[1],
            "project_status": project_result[3],
            "project_age_days": (datetime.now() - project_result[4]).days,
            "business_context": business_context,
            "recent_conversation_context": recent_context,
            "documents_created": documents_created,
            "total_interactions": len(recent_chats),
            "context_summary": generate_context_summary(business_context, recent_context)
        }
        
    except Exception as e:
        print(f"❌ Error getting project context: {e}")
        return {}

def generate_context_summary(business_context, recent_chats):
    """Genera un resumen del contexto para incluir en prompts"""
    try:
        summary_parts = []
        
        # Información del negocio
        if business_context.get("startup_name"):
            summary_parts.append(f"Startup: {business_context['startup_name']}")
        
        if business_context.get("industry"):
            summary_parts.append(f"Industria: {business_context['industry']}")
        
        if business_context.get("stage"):
            summary_parts.append(f"Etapa: {business_context['stage']}")
        
        if business_context.get("business_type"):
            summary_parts.append(f"Tipo: {business_context['business_type']}")
        
        if business_context.get("problem_solving"):
            summary_parts.append(f"Problema que resuelve: {business_context['problem_solving']}")
        
        # Contexto reciente
        if recent_chats:
            summary_parts.append(f"Últimas conversaciones: {len(recent_chats)} interacciones recientes")
        
        return " | ".join(summary_parts) if summary_parts else "Nuevo proyecto sin contexto previo"
        
    except Exception as e:
        print(f"❌ Error generating context summary: {e}")
        return "Error generando resumen de contexto"

def detect_investor_search_intent(user_message, user_language='en'):
    """Detecta si el usuario quiere buscar inversores ESPECÍFICOS usando Gemini - VERSIÓN CORREGIDA"""
    try:
        detection_prompts = {
            'es': f"""
            Analiza si el usuario quiere BUSCAR INVERSORES ESPECÍFICOS para su startup:
            "{user_message}"
            
            Responde SOLO 'true' o 'false'.
            
            ✅ SÍ es búsqueda de inversores (responde 'true'):
            - "Buscar inversores para fintech"
            - "Encuentrame VCs de Madrid" 
            - "Necesito inversores"
            - "Startup Profile for Investor Matching"
            - Cualquier descripción que incluya: "funding", "Series A/B/C", "seed", "€", "$", "million", "ARR"
            - Descripciones detalladas de startups con métricas o industria
            - "I need investors for my..."
            - Texto estructurado con información de startup (sector, modelo, etc)
            
            ❌ NO es búsqueda (responde 'false'):
            - "Genera un pitch deck"
            - "Crea un documento"
            - "Escribe un plan"
            - "Cómo convencer inversores" (consejo, no búsqueda)
            - Preguntas teóricas sobre inversión
            """,
            
            'en': f"""
            Analyze if user wants to SEARCH FOR SPECIFIC INVESTORS for their startup:
            "{user_message}"
            
            Respond ONLY with 'true' or 'false'.
            
            ✅ YES it's investor search (respond 'true'):
            - "Search for fintech investors"
            - "Find VCs in London"
            - "I need investors"
            - "Startup Profile for Investor Matching"
            - Any description including: "funding", "Series A/B/C", "seed", "€", "$", "million", "ARR"
            - Detailed startup descriptions with metrics or industry
            - "I need investors for my..."
            - Structured text with startup info (sector, model, etc)
            
            ❌ NO it's not search (respond 'false'):
            - "Generate a pitch deck"
            - "Create a document"
            - "Write a plan"
            - "How to convince investors" (advice, not search)
            - Theoretical questions about investment
            """
        }
        
        prompt = detection_prompts.get(user_language, detection_prompts['en'])
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.1,
                "max_output_tokens": 10,
            }
        )
        
        result = response.text.strip().lower() == 'true'
        
        # FALLBACK MEJORADO: Si Gemini falla, buscar patrones específicos
        if not result:
            # Patrones que SIEMPRE indican búsqueda de inversores
            strong_investor_patterns = [
                'startup profile for investor',
                'need investors for',
                'looking for.*funding',
                'series [a-c]',
                'seeking.*investment',
                '\\d+.*million.*arr',
                'funding.*needed',
                'investor matching'
            ]
            
            message_lower = user_message.lower()
            for pattern in strong_investor_patterns:
                if re.search(pattern, message_lower):
                    result = True
                    print(f"🔄 Fallback pattern matched: {pattern}")
                    break
        
        print(f"🤖 Intent detection: '{user_message[:50]}...' → {result}")
        return result
        
    except Exception as e:
        print(f"❌ Error detecting investor search intent: {e}")
        # Fallback más agresivo
        investor_keywords = ['investor', 'funding', 'series', 'seed', 'vc', 'capital']
        message_lower = user_message.lower()
        return any(keyword in message_lower for keyword in investor_keywords)

def detect_document_generation_intent(user_message, user_language='en'):
    """Detecta si el usuario quiere generar un documento profesional - VERSIÓN MEJORADA"""
    try:
        # IMPORTANTE: Si ya menciona inversores, NO es documento
        if 'investor' in user_message.lower() or 'funding' in user_message.lower():
            return False
            
        detection_prompts = {
            'es': f"""
            Analiza si el usuario quiere que GENERES un DOCUMENTO profesional:
            "{user_message}"
            
            Responde SOLO 'true' o 'false'.
            
            ✅ SÍ quiere documento (responde 'true'):
            - "Hazme un pitch deck"
            - "Genera un plan de negocios"
            - "Crea un plan de marketing"
            - "Redacta el modelo financiero"
            - Uso explícito de verbos: "genera", "crea", "hazme", "escribe"
            
            ❌ NO es documento (responde 'false'):
            - Cualquier mención de "inversores" o "funding"
            - "Startup Profile for Investor Matching"
            - Descripciones de negocio sin pedir documento
            - Preguntas generales
            """,
            
            'en': f"""
            Analyze if user wants you to GENERATE a professional DOCUMENT:
            "{user_message}"
            
            Respond ONLY 'true' or 'false'.
            
            ✅ YES wants document (respond 'true'):
            - "Make me a pitch deck"
            - "Generate a business plan"
            - "Create a marketing plan"
            - "Write the financial model"
            - Explicit use of verbs: "generate", "create", "make", "write"
            
            ❌ NO document (respond 'false'):
            - Any mention of "investors" or "funding"
            - "Startup Profile for Investor Matching"
            - Business descriptions without asking for document
            - General questions
            """
        }
        
        prompt = detection_prompts.get(user_language, detection_prompts['en'])
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            prompt,
            generation_config={"temperature": 0.1, "max_output_tokens": 10}
        )
        
        return response.text.strip().lower() == 'true'
        
    except Exception as e:
        print(f"❌ Error detecting document intent: {e}")
        return False

def identify_document_type(message, language='en'):
    """Identifica qué tipo de documento quiere el usuario"""
    message_lower = message.lower()
    
    # Mapeo de palabras clave a tipos de documento
    mappings = {
        'pitch_deck': ['pitch', 'deck', 'investor deck', 'pitch deck'],
        'business_plan': ['business plan', 'plan de negocios', 'plan de negocio'],
        'marketing_plan': ['marketing', 'plan de marketing', 'growth plan'],
        'financial_model': ['financial', 'modelo financiero', 'proyecciones', 'projections']
    }
    
    for doc_type, keywords in mappings.items():
        if any(keyword in message_lower for keyword in keywords):
            return doc_type
    
    return None

def generate_professional_document_in_chat(user_id, project_id, document_type, user_context, user_language):
    """Genera documento profesional dentro del chat"""
    try:
        # Costo en créditos
        doc_costs = {
            'pitch_deck': 50,
            'business_plan': 75,
            'marketing_plan': 60,
            'financial_model': 70
        }
        
        credits_required = doc_costs.get(document_type, 50)
        
        # Verificar créditos
        if not has_sufficient_credits(user_id, credits_required):
            return {
                'error': 'insufficient_credits',
                'required': credits_required,
                'available': get_user_credits(user_id)
            }
        
        # Contexto del negocio
        business_context = user_context.get('business_context', {})
        
        # Prompts profesionales por tipo de documento
        doc_prompts = {
            'pitch_deck': f"""
            Create a PROFESSIONAL pitch deck for {business_context.get('startup_name', 'this startup')}.
            
            Industry: {business_context.get('industry', 'Technology')}
            Stage: {business_context.get('stage', 'Seed')}
            Problem: {business_context.get('problem_solving', 'To be defined')}
            
            STRUCTURE (each slide with ## Title):
            
            ## 1. Cover Slide
            [Company Name]
            [Powerful tagline that captures the essence]
            [Contact information]
            
            ## 2. The Problem
            - Specific pain point with real data
            - Who suffers from this (TAM)
            - Current inadequate solutions
            - Cost of the problem
            
            ## 3. Our Solution
            - Clear value proposition
            - How it works (simple explanation)
            - Key differentiators
            - Why now?
            
            ## 4. Market Opportunity
            - TAM: $X billion (with source)
            - SAM: $X billion 
            - SOM: $X million (realistic first 3 years)
            - Growth rate and trends
            
            ## 5. Product Demo
            - Key features (3-5 max)
            - User journey
            - Screenshots/mockups description
            - Tech stack highlights
            
            ## 6. Business Model
            - Revenue streams
            - Pricing strategy
            - Unit economics
            - Path to profitability
            
            ## 7. Go-to-Market Strategy
            - Customer acquisition channels
            - Sales strategy
            - Marketing approach
            - Key partnerships
            
            ## 8. Competition
            - Competitive landscape matrix
            - Our unique advantages
            - Barriers to entry
            - Network effects
            
            ## 9. Traction
            - Current metrics (users, revenue, growth)
            - Key achievements
            - Customer testimonials
            - Notable partnerships
            
            ## 10. Team
            - Founders (background, expertise)
            - Key team members
            - Advisors
            - Why we'll win
            
            ## 11. Financial Projections
            - 3-year revenue projection
            - Key metrics evolution
            - Burn rate and runway
            - Path to break-even
            
            ## 12. The Ask
            - Funding amount: $X
            - Use of funds (specific allocation)
            - Milestones to achieve
            - Expected valuation
            
            ## 13. Vision
            - 5-year vision
            - Impact on the market
            - Exit strategy potential
            
            Make it compelling, data-driven, and investor-ready. Use specific numbers and avoid generic statements.
            """
        }
        
        # Generar con Gemini
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            doc_prompts.get(document_type, doc_prompts['pitch_deck']),
            generation_config={
                "temperature": 0.7,
                "max_output_tokens": 6000
            }
        )
        
        # Cobrar créditos
        charge_credits(user_id, credits_required)
        
        # Respuesta combinada
        intro_messages = {
            'es': f"✅ **¡{document_type.replace('_', ' ').title()} profesional generado!**\n\nAquí está tu documento listo para usar:\n\n---\n\n",
            'en': f"✅ **Professional {document_type.replace('_', ' ').title()} generated!**\n\nHere's your document ready to use:\n\n---\n\n"
        }
        
        final_response = intro_messages.get(user_language, intro_messages['en']) + response.text
        
        # Guardar en historial
        save_neural_interaction(user_id, {
            'bot': 'document_generator',
            'input': f"Generate {document_type}",
            'response': final_response,
            'credits_used': credits_required,
            'context': user_context,
            'session_id': user_context.get('session_id'),
            'project_id': project_id,
            'document_type': document_type
        })
        
        return {
            'success': True,
            'bot': 'interactive_mentor',
            'response': final_response,
            'credits_charged': credits_required,
            'credits_remaining': get_user_credits(user_id),
            'document_generated': True,
            'document_type': document_type
        }
        
    except Exception as e:
        print(f"❌ Error generating document: {e}")
        return {
            'error': 'Could not generate document',
            'details': str(e)
        }

# ==============================================================================
#           BOT SYSTEM
# ==============================================================================

# ==============================================================================
#           BOT SYSTEM - VERSIÓN CORREGIDA
# ==============================================================================

class BotManager:
    def process_user_request(self, user_input, user_context, user_id):
        """Procesa request del usuario con tracking correcto de créditos"""
        try:
            # 1. VERIFICAR CRÉDITOS ACTUALES
            credits_before = get_user_credits(user_id)
            required_credits = CREDIT_COSTS.get('basic_bot', 5)
            
            print(f"🤖 Bot procesando:")
            print(f"   - Usuario: {user_id}")
            print(f"   - Créditos antes: {credits_before}")
            print(f"   - Créditos requeridos: {required_credits}")
            print(f"   - Idioma detectado: {user_context.get('user_language', 'en')}")
            
            if credits_before < required_credits:
                return {
                    'error': 'Insufficient credits',
                    'required': required_credits,
                    'available': credits_before
                }
            
            # 2. GENERAR RESPUESTA CON GEMINI
            prompt = self._build_smart_prompt(user_input, user_context)
            
            model = genai.GenerativeModel(MODEL_NAME)
            ai_response = model.generate_content(
                prompt,
                generation_config={
                    "temperature": 0.7,
                    "top_p": 0.95,
                    "top_k": 40,
                    "max_output_tokens": 2000,
                }
            )
            
            # 3. COBRAR CRÉDITOS
            print(f"💳 Cobrando {required_credits} créditos...")
            credits_after_charge = charge_credits(user_id, required_credits)
            
            if credits_after_charge is None:
                return {'error': 'Could not charge credits'}
            
            print(f"✅ Créditos después del cobro: {credits_after_charge}")
            
            # 4. GUARDAR INTERACCIÓN
            save_neural_interaction(user_id, {
                'bot': 'interactive_mentor',
                'input': user_input,
                'response': ai_response.text,
                'credits_used': required_credits,
                'context': user_context,
                'session_id': user_context.get('session_id'),
                'project_id': user_context.get('project_id')
            })
            
            # 5. RETORNAR RESPUESTA
            return {
                'bot': 'interactive_mentor',
                'response': ai_response.text,
                'credits_charged_by_bot': required_credits,
                'processing_success': True
            }
            
        except Exception as e:
            print(f"❌ Error in bot processing: {e}")
            import traceback
            traceback.print_exc()
            return {'error': f'Bot error: {str(e)}'}

    def _build_smart_prompt(self, user_input, user_context):
        """Prompt inteligente para respuestas naturales y contextuales"""
        business_context = user_context.get('business_context', {})
        user_plan = user_context.get('user_plan', 'free')
        user_language = user_context.get('user_language', 'en')
        
        # Detectar longitud de input para ajustar respuesta
        input_length = len(user_input.split())
        is_greeting = any(greeting in user_input.lower() for greeting in ['hi', 'hello', 'hey', 'hola'])
        
        language_instructions = {
            'es': "Responde SIEMPRE en español.",
            'en': "Respond ALWAYS in English."
        }.get(user_language, "Respond in the same language as the user.")
        
        # Construir sección de roadmap si está disponible
        roadmap_section = ""
        if user_context.get('show_roadmap') and user_context.get('roadmap'):
            roadmap = user_context['roadmap']
            roadmap_section = f"""
    
    CURRENT ROADMAP FOR USER:
    - Stage: {roadmap.get('current', '')}
    - Urgent: {roadmap.get('urgent_action', '')}
    - CTA to include naturally: {roadmap.get('cta', '')}
    """
        
        prompt = f"""
        You are an expert startup mentor who has helped 200+ startups raise capital.
        
        {language_instructions}
        
        USER CONTEXT:
        - User plan: {user_plan} (this is their subscription level, NOT their target users)
        - Business: {business_context.get('startup_name', 'Not defined yet')}
        - Stage: {business_context.get('stage', 'Idea')}
        - Industry: {business_context.get('industry', 'Not specified')}
        {roadmap_section}
        
        USER MESSAGE: {user_input}
        
        RESPONSE RULES:
        1. Match response length to input:
           - Greeting (hi/hello): 1-2 sentences friendly response
           - Short question: 2-3 paragraphs max
           - Complex question: Detailed but structured response
        
        2. Be a friend AND mentor:
           - Casual and approachable
           - But professional when needed
           - Use examples only when relevant
        
        3. Guide towards action:
           - If they have just an idea, suggest next steps
           - Offer to create documents naturally: "Need a pitch deck for that?"
           - Don't overwhelm with all options at once
           - If roadmap CTA is provided, work it into your response naturally
        
        4. NEVER mention their subscription plan unless they ask
        5. Keep first response under 150 words if input is under 10 words
        
        Respond naturally:
        """
        
        return prompt

    def get_startup_roadmap(self, user_context):
        """Genera roadmap INTELIGENTE que guía + hace upsell según tipo de usuario"""
        stage = user_context.get('business_context', {}).get('stage', 'idea')
        user_plan = user_context.get('user_plan', 'free')
        user_language = user_context.get('user_language', 'en')
        
        # Roadmaps personalizados por TIPO DE USUARIO + PLAN
        roadmaps = {
            'idea': {
                'es': {
                    'current': '💡 Fase Idea',
                    'urgent_action': '⚡ ACCIÓN INMEDIATA: Valida tu idea en 48h',
                    'free_steps': [
                        '✅ Define el problema (Chat conmigo)',
                        '✅ Identifica 10 early adopters',
                        '❌ Buscar inversores angel (Requiere Growth)',
                        '❌ Generar pitch deck profesional (50 créditos)'
                    ],
                    'growth_unlock': '🚀 Con Growth Plan: Encuentra 20+ inversores angel perfectos para ideas',
                    'cta': '¿Quieres que busque inversores angel para tu idea ahora?'
                },
                'en': {
                    'current': '💡 Idea Stage',
                    'urgent_action': '⚡ IMMEDIATE ACTION: Validate your idea in 48h',
                    'free_steps': [
                        '✅ Define the problem (Chat with me)',
                        '✅ Identify 10 early adopters',
                        '❌ Find angel investors (Requires Growth)',
                        '❌ Generate professional pitch deck (50 credits)'
                    ],
                    'growth_unlock': '🚀 With Growth Plan: Find 20+ perfect angel investors for ideas',
                    'cta': 'Want me to find angel investors for your idea now?'
                }
            },
            'mvp': {
                'es': {
                    'current': '🚀 Fase MVP',
                    'urgent_action': '⚡ CRÍTICO: Necesitas tracción ANTES de buscar inversión',
                    'free_steps': [
                        '✅ Conseguir primeros 100 usuarios',
                        '✅ Medir retención semanal',
                        '❌ Buscar VCs seed stage (Requiere Growth)',
                        '❌ Plan financiero profesional (70 créditos)'
                    ],
                    'growth_unlock': '🎯 Con Growth: 20+ VCs que invierten en MVPs sin revenue',
                    'cta': 'Busco VCs que inviertan en MVPs como el tuyo. ¿Empezamos?'
                },
                'en': {
                    'current': '🚀 MVP Stage',
                    'urgent_action': '⚡ CRITICAL: You need traction BEFORE seeking investment',
                    'free_steps': [
                        '✅ Get first 100 users',
                        '✅ Measure weekly retention',
                        '❌ Find seed stage VCs (Requires Growth)',
                        '❌ Professional financial model (70 credits)'
                    ],
                    'growth_unlock': '🎯 With Growth: 20+ VCs that invest in pre-revenue MVPs',
                    'cta': 'I can find VCs that invest in MVPs like yours. Start now?'
                }
            },
            'seed': {
                'es': {
                    'current': '💰 Fase Seed',
                    'urgent_action': '⚡ MOMENTO CLAVE: Series A en 12-18 meses',
                    'free_steps': [
                        '✅ Escalar a €1M ARR',
                        '✅ Contratar primeros 5 empleados',
                        '❌ Buscar fondos Series A (Requiere Growth)',
                        '❌ Buscar CTOs/CMOs (Requiere Growth)'
                    ],
                    'growth_unlock': '💎 Con Growth: Encuentra Series A + Empleados clave de fondos top',
                    'cta': '¿Necesitas Series A o un CTO? Puedo buscar ambos.'
                },
                'en': {
                    'current': '💰 Seed Stage',
                    'urgent_action': '⚡ KEY MOMENT: Series A in 12-18 months',
                    'free_steps': [
                        '✅ Scale to $1M ARR',
                        '✅ Hire first 5 employees',
                        '❌ Find Series A funds (Requires Growth)',
                        '❌ Find CTOs/CMOs (Requires Growth)'
                    ],
                    'growth_unlock': '💎 With Growth: Find Series A + Key employees from top funds',
                    'cta': 'Need Series A or a CTO? I can find both.'
                }
            },
            'corporate': {
                'es': {
                    'current': '🏢 Corporación Establecida',
                    'urgent_action': '⚡ OPORTUNIDAD: Innovación corporativa',
                    'free_steps': [
                        '✅ Identificar nuevas verticales',
                        '✅ Estrategia de innovación',
                        '❌ Corporate VCs estratégicos (Pro Plan)',
                        '❌ M&A targets automático (Pro Plan)'
                    ],
                    'growth_unlock': '🎯 Con Pro: Outreach automatizado a partners estratégicos',
                    'cta': '¿Buscamos corporate VCs o targets de adquisición?'
                },
                'en': {
                    'current': '🏢 Established Corporation',
                    'urgent_action': '⚡ OPPORTUNITY: Corporate innovation',
                    'free_steps': [
                        '✅ Identify new verticals',
                        '✅ Innovation strategy',
                        '❌ Strategic corporate VCs (Pro Plan)',
                        '❌ Automated M&A targets (Pro Plan)'
                    ],
                    'growth_unlock': '🎯 With Pro: Automated outreach to strategic partners',
                    'cta': 'Should we find corporate VCs or acquisition targets?'
                }
            }
        }
        
        # Seleccionar roadmap según stage
        roadmap_key = 'corporate' if stage == 'series_b_plus' else stage
        if roadmap_key not in roadmaps:
            roadmap_key = 'idea'
            
        roadmap = roadmaps[roadmap_key].get(user_language, roadmaps[roadmap_key]['en'])
        
        # Si es usuario FREE, añadir urgencia
        if user_plan == 'free':
            roadmap['upgrade_urgency'] = {
                'es': '⏰ Usuarios Growth encuentran inversores 3x más rápido',
                'en': '⏰ Growth users find investors 3x faster'
            }.get(user_language, '⏰ Growth users find investors 3x faster')
        
        return roadmap

    def _format_recent_conversations(self, conversations):
        """Formatea conversaciones recientes para contexto"""
        if not conversations:
            return "No previous conversations"
        
        formatted = []
        for conv in conversations[-3:]:  # Últimas 3
            formatted.append(f"- User said: {conv.get('user_said', '')}")
            formatted.append(f"- Assistant responded: {conv.get('assistant_responded', '')}")
        
        return '\n'.join(formatted)

# Actualizar la instancia global
bot_manager = BotManager()

class InvestorSearchSimple:
    """
    Búsqueda inteligente de inversores - VERSIÓN MEJORADA CON GEMINI
    """
    
    def __init__(self, engine):
        self.engine = engine
        self.vectorizer = TfidfVectorizer(stop_words='english', max_features=500)
        # Cargar contextos de los archivos .txt
        self.ubicaciones_context = self._load_context_file('ubicaciones.txt')
        self.etapas_context = self._load_context_file('etapas.txt')
        self.categorias_context = self._load_context_file('categorias.txt')
        print("✅ Motor de búsqueda mejorado inicializado")
    
    def _load_context_file(self, filepath):
        """Carga archivo de contexto"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"⚠️ No se pudo cargar {filepath}: {e}")
            return ""
    
    def buscar_inversores(self, query):
        """FUNCIÓN PRINCIPAL: Busca inversores usando Gemini para keywords inteligentes"""
        try:
            print(f"🔍 Búsqueda inteligente: '{query}'")
            
            # 1. Cargar inversores
            inversores = self._cargar_inversores()
            if inversores.empty:
                return {"error": "No hay inversores en la base de datos"}
            
            # 2. EXTRAER KEYWORDS INTELIGENTES CON GEMINI
            keywords_gemini = self._get_keywords_from_gemini_v2(query)
            if not keywords_gemini:
                return {"error": "No se pudieron extraer keywords"}
            
            print(f"🧠 Keywords extraídas: {keywords_gemini}")
            
            # 3. Calcular puntuaciones con PESOS AJUSTADOS
            con_puntuacion = self._calculate_match_score_v2(inversores, keywords_gemini)
            
            # 4. Top 21 con scores inventados
            mejores_21 = con_puntuacion.head(21)
            
            # 5. Scores realistas 72-99
            import random
            for idx in mejores_21.index:
                position = list(mejores_21.index).index(idx)
                if position < 5:
                    fake_score = random.uniform(88, 99)
                elif position < 10:
                    fake_score = random.uniform(80, 92)
                else:
                    fake_score = random.uniform(72, 85)
                mejores_21.loc[idx, 'score_final'] = fake_score
            
            mejores_21 = mejores_21.sort_values('score_final', ascending=False)
            
            return self._formatear_resultados(mejores_21, query)
            
        except Exception as e:
            print(f"❌ Error en búsqueda: {e}")
            return {"error": f"Búsqueda falló: {str(e)}"}
    
    def _get_keywords_from_gemini_v2(self, query):
        """Extrae keywords inteligentes usando Gemini - VERSIÓN MEJORADA"""
        try:
            model = genai.GenerativeModel(MODEL_NAME)
            prompt = f"""
            Analiza esta búsqueda de inversores y extrae keywords PRECISAS y EXPANDIDAS.
            
            Consulta: "{query}"
            
            Instrucciones:
            1. Identifica ubicaciones, etapas de inversión y categorías/industrias
            2. Usa ÚNICAMENTE los términos de estos contextos para estandarizar
            3. Genera 5-10 keywords expandidas ALTAMENTE RELEVANTES para cada categoría
            4. NO inventes términos que no estén relacionados
            
            Contexto Ubicaciones (ejemplos):
            {self.ubicaciones_context[:500]}...
            
            Contexto Etapas (ejemplos):
            {self.etapas_context[:500]}...
            
            Contexto Categorías (ejemplos):
            {self.categorias_context[:500]}...
            
            Devuelve SOLO este JSON exacto:
            {{
                "ubicacion": {{
                    "primary": ["ubicaciones exactas mencionadas"],
                    "expanded": ["ubicaciones relacionadas del contexto"]
                }},
                "etapa": {{
                    "primary": ["etapas exactas mencionadas"],
                    "expanded": ["etapas relacionadas del contexto"]
                }},
                "categoria": {{
                    "primary": ["categorías exactas mencionadas"],
                    "expanded": ["categorías relacionadas del contexto"]
                }}
            }}
            """
            
            response = model.generate_content(
                prompt,
                generation_config={"temperature": 0.2, "max_output_tokens": 1000}
            )
            
            # Extraer JSON
            response_text = response.text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:-3]
            
            keywords = json.loads(response_text)
            
            # Limpiar y normalizar
            for cat in keywords:
                for subcat in ['primary', 'expanded']:
                    if subcat in keywords[cat]:
                        keywords[cat][subcat] = [k.strip().lower() for k in keywords[cat][subcat]]
            
            return keywords
            
        except Exception as e:
            print(f"❌ Error Gemini: {e}")
            return None
    
    def _calculate_match_score_v2(self, df, keywords):
        """Calcula scores con PESOS REGULATORIOS (ubicación 45%)"""
        try:
            # Preparar datos de inversores
            df['ubicacion_list'] = df['ubicacion'].str.lower().str.split(',')
            df['etapas_list'] = df['etapas'].str.lower().str.split(',')
            df['categorias_list'] = df['categorias'].str.lower().str.split(',')
            
            # PESOS AJUSTADOS - Ubicación más importante por regulaciones
            WEIGHTS = {
                'ubicacion': 0.334,  # 45% por regulaciones
                'etapa': 0.333,      # 30% 
                'categoria': 0.333   # 25%
            }
            
            scores = []
            for _, inversor in df.iterrows():
                score = 0
                
                # Ubicación
                ubi_primary = set(keywords['ubicacion']['primary'])
                ubi_expanded = set(keywords['ubicacion']['expanded'])
                ubi_inversor = set([u.strip() for u in inversor['ubicacion_list'] if isinstance(u, str)])
                
                ubi_match_primary = len(ubi_primary.intersection(ubi_inversor))
                ubi_match_expanded = len(ubi_expanded.intersection(ubi_inversor))
                score += (ubi_match_primary * 3 + ubi_match_expanded) * WEIGHTS['ubicacion']
                
                # Etapa
                eta_primary = set(keywords['etapa']['primary'])
                eta_expanded = set(keywords['etapa']['expanded'])
                eta_inversor = set([e.strip() for e in inversor['etapas_list'] if isinstance(e, str)])
                
                eta_match_primary = len(eta_primary.intersection(eta_inversor))
                eta_match_expanded = len(eta_expanded.intersection(eta_inversor))
                score += (eta_match_primary * 3 + eta_match_expanded) * WEIGHTS['etapa']
                
                # Categoría
                cat_primary = set(keywords['categoria']['primary'])
                cat_expanded = set(keywords['categoria']['expanded'])
                cat_inversor = set([c.strip() for c in inversor['categorias_list'] if isinstance(c, str)])
                
                cat_match_primary = len(cat_primary.intersection(cat_inversor))
                cat_match_expanded = len(cat_expanded.intersection(cat_inversor))
                score += (cat_match_primary * 3 + cat_match_expanded) * WEIGHTS['categoria']
                
                scores.append(score)
            
            df['score_final'] = scores
            return df.sort_values('score_final', ascending=False)
            
        except Exception as e:
            print(f"❌ Error calculando scores: {e}")
            return df.assign(score_final=50)
    
    def _cargar_inversores(self):
        """Carga inversores de la BD"""
        try:
            query = """
                SELECT 
                    id,
                    "Company_Name" as nombre,
                    "Company_Description" as descripcion,
                    "Company_Location" as ubicacion,
                    "Investing_Stage" as etapas,
                    "Investment_Categories" as categorias,
                    "Company_Linkedin" as linkedin
                FROM investors
                WHERE "Company_Name" IS NOT NULL 
                LIMIT 1000
            """
            
            df = pd.read_sql(query, self.engine)
            df = df.fillna('')
            
            df['texto_busqueda'] = (
                df['nombre'] + ' ' +
                df['descripcion'] + ' ' +
                df['ubicacion'] + ' ' +
                df['etapas'] + ' ' +
                df['categorias']
            ).str.lower()
            
            return df
            
        except Exception as e:
            print(f"❌ Error cargando inversores: {e}")
            return pd.DataFrame()
    
    def _formatear_resultados(self, df, query):
        """Formatea resultados finales"""
        resultados = []
        
        for i, (_, inversor) in enumerate(df.iterrows()):
            resultado = {
                'investor_id': str(inversor.get('id', '')),
                'company_name': inversor.get('nombre', ''),
                'description': inversor.get('descripcion', '')[:300] + '...' if len(str(inversor.get('descripcion', ''))) > 300 else inversor.get('descripcion', ''),
                'location': inversor.get('ubicacion', ''),
                'investing_stages': inversor.get('etapas', ''),
                'investment_categories': inversor.get('categorias', ''),
                'linkedin_url': inversor.get('linkedin', ''),
                'match_score': inversor.get('score_final', 75)
            }
            resultados.append(resultado)
        
        return {
            'search_type': 'gemini_intelligent_v3',
            'query': query,
            'results': resultados,
            'total_found': len(resultados),
            'max_results': 21,
            'success': True
        }
                     
def save_investors_to_project_simple(user_id, project_id, investors_data):
    """
    Guarda los 20 mejores inversores en project_saved_investors
    """
    try:
        saved_count = 0
        
        with engine.connect() as conn:
            for investor in investors_data:
                investor_id = investor['investor_id']
                
                # Verificar si ya existe esta relación
                existing = conn.execute(
                    text("""
                        SELECT 1 FROM project_saved_investors 
                        WHERE project_id = :project_id AND investor_id = :investor_id
                    """),
                    {"project_id": project_id, "investor_id": investor_id}
                ).fetchone()
                
                if not existing:
                    # Insertar nueva relación
                    conn.execute(
                        text("""
                            INSERT INTO project_saved_investors (project_id, investor_id, added_at) 
                            VALUES (:project_id, :investor_id, NOW())
                        """),
                        {"project_id": project_id, "investor_id": investor_id}
                    )
                    saved_count += 1
                    print(f"✅ Saved investor: {investor.get('company_name', investor_id)}")
            
            conn.commit()
            
        print(f"✅ Total investors saved: {saved_count}")
        return saved_count
        
    except Exception as e:
        print(f"❌ Error saving investors: {e}")
        import traceback
        traceback.print_exc()
        return 0

def save_search_to_history_simple(user_id, query, results_count, credits_spent):
    """
    Guarda búsqueda en search_history para trackear lo que busca el usuario
    """
    try:
        with engine.connect() as conn:
            conn.execute(
                text("""
                    INSERT INTO search_history (
                        id, user_id, search_type, query, results_count, 
                        credits_spent, filters_used, created_at
                    ) VALUES (
                        :id, :user_id, 'investors', :query, :results_count,
                        :credits_spent, :filters_used, NOW()
                    )
                """),
                {
                    "id": str(uuid.uuid4()),
                    "user_id": user_id,
                    "query": query,
                    "results_count": results_count,
                    "credits_spent": credits_spent,
                    "filters_used": json.dumps({
                        "executed_from": "chat",
                        "auto_search": True,
                        "max_results": 20
                    })
                }
            )
            conn.commit()
            print(f"✅ Search saved to history: {query}")
            
    except Exception as e:
        print(f"❌ Error saving search history: {e}")

def create_investors_table_for_chat(investors, language='en'):
    """
    Crea tabla bonita de inversores para mostrar EN EL CHAT
    """
    try:
        if not investors:
            return "No investors to display"
        
        if language == 'es':
            headers = ["#", "Fondo de Inversión", "Ubicación", "Etapas", "Categorías", "Score"]
        else:
            headers = ["#", "Investment Fund", "Location", "Stages", "Categories", "Score"]
        
        # Crear tabla markdown
        table_lines = []
        table_lines.append("| " + " | ".join(headers) + " |")
        table_lines.append("|" + "|".join([":---:"] + ["---"] * (len(headers)-1)) + "|")
        
        for i, investor in enumerate(investors, 1):
            # Truncar textos largos para que se vea bien en el chat
            company_name = investor.get('company_name', 'N/A')
            if len(company_name) > 25:
                company_name = company_name[:22] + "..."
                
            location = investor.get('location', 'N/A')
            if len(location) > 15:
                location = location[:12] + "..."
                
            stages = investor.get('investing_stages', 'N/A')
            if len(stages) > 15:
                stages = stages[:12] + "..."
                
            categories = investor.get('investment_categories', 'N/A')
            if len(categories) > 20:
                categories = categories[:17] + "..."
                
            score = f"{investor.get('match_score', 0):.0f}%"
            
            row = [
                f"**{i}**",
                f"**{company_name}**",
                location,
                stages,
                categories,
                f"**{score}**"
            ]
            
            table_lines.append("| " + " | ".join(row) + " |")
            
            # Solo mostrar primeros 10 en el chat
            if i >= 21:
                break
        
        # Si hay más de 10, agregar nota
        if len(investors) < 10:
            if language == 'es':
                table_lines.append(f"\n*...y {len(investors) - 10} inversores más guardados en tu página de Inversores*")
            else:
                table_lines.append(f"\n*...and {len(investors) - 10} more investors saved to your Investors page*")
        
        return "\n".join(table_lines)
        
    except Exception as e:
        print(f"❌ Error creating table: {e}")
        return "Error creating investors table"
                   
# ==============================================================================
#           ROUTES
# ==============================================================================

@app.route('/')
def home():
    """Home endpoint"""
    return jsonify({
        'status': 'online',
        'version': '2.0.0',
        'message': '🚀 0Bullshit Backend API - CORS ARREGLADO DEFINITIVAMENTE! 🎉',
        'auth_methods': ['manual', 'google'],
        'database_connected': engine is not None,
        'google_auth_enabled': GOOGLE_CLIENT_ID is not None,
        'cors_status': '✅ CORS KILLER ACTIVADO',
        'timestamp': datetime.now().isoformat()
    })

@app.route('/cors-test', methods=['GET', 'POST', 'OPTIONS'])
def cors_test():
    """Endpoint para testear que CORS funciona - CORS KILLER TEST"""
    return jsonify({
        'message': '🎉 CORS está funcionando PERFECTAMENTE!',
        'method': request.method,
        'origin': request.headers.get('Origin', 'No origin'),
        'user_agent': request.headers.get('User-Agent', 'No user agent')[:50] + '...',
        'headers_received': dict(request.headers),
        'timestamp': datetime.now().isoformat(),
        'cors_status': '✅ CORS KILLER FUNCIONANDO',
        'backend_status': 'ONLINE',
        'database_status': 'CONNECTED' if engine else 'DISCONNECTED'
    })

@app.route('/health')
def health_check():
    """Detailed health check"""
    try:
        # Test database connection
        if engine:
            with engine.connect() as conn:
                conn.execute(text("SELECT 1"))
            db_status = "connected"
        else:
            db_status = "disconnected"
        
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.utcnow().isoformat(),
            'database': db_status,
            'google_auth': 'enabled' if GOOGLE_CLIENT_ID else 'disabled',
            'gemini_api': 'enabled' if GEMINI_API_KEY else 'disabled',
            'environment': 'production' if 'render.com' in os.environ.get('RENDER_EXTERNAL_URL', '') else 'development',
            'cors_status': '✅ CORS KILLER ACTIVE'
        })
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e)
        }), 500

@app.route('/auth/test', methods=['GET'])
def test_auth():
    """Test endpoint para verificar que el sistema funciona"""
    try:
        return jsonify({
            'status': 'ok',
            'message': 'Auth system is working',
            'database_connected': engine is not None,
            'timestamp': datetime.now().isoformat(),
            'cors_working': True
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/auth/register', methods=['POST'])
def register():
    """Registro de usuario - CON DETECCIÓN DE IDIOMA"""
    try:
        print("🚀 Register endpoint called")
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        print(f"📥 Received data: {data}")
        
        # Detectar idioma del formulario (basado en el campo first_name o cualquier texto)
        detected_language = detect_user_language(data.get('first_name', '') + ' ' + data.get('last_name', ''))
        
        # Mensajes de error multiidioma
        error_messages = {
            'es': {
                'missing_field': 'Campo requerido faltante: {}',
                'invalid_email': 'Formato de email inválido',
                'email_registered': 'Email ya registrado',
                'error_creating': 'Error creando usuario',
                'error_token': 'Error generando token',
                'success': 'Usuario registrado exitosamente'
            },
            'en': {
                'missing_field': 'Missing required field: {}',
                'invalid_email': 'Invalid email format',
                'email_registered': 'Email already registered',
                'error_creating': 'Error creating user',
                'error_token': 'Error generating token',
                'success': 'User registered successfully'
            }
        }
        
        msgs = error_messages.get(detected_language, error_messages['en'])
        
        # Validar campos requeridos
        required_fields = ['email', 'password', 'first_name', 'last_name']
        for field in required_fields:
            if field not in data or not data[field]:
                return jsonify({'error': msgs['missing_field'].format(field)}), 400
        
        # Validar formato de email
        email = data['email'].lower().strip()
        if '@' not in email or '.' not in email:
            return jsonify({'error': msgs['invalid_email']}), 400
        
        # Validar contraseña
        password = data['password']
        is_valid, message = validate_password_strength(password)
        if not is_valid:
            return jsonify({'error': message}), 400
        
        # Verificar si email ya existe
        existing_user = get_user_by_email(email)
        if existing_user:
            return jsonify({'error': msgs['email_registered']}), 400
        
        # Crear usuario
        user_id = create_user(
            email=email,
            password=password,
            first_name=data['first_name'].strip(),
            last_name=data['last_name'].strip()
        )
        
        if not user_id:
            return jsonify({'error': msgs['error_creating']}), 500
        
        # Inicializar memoria neuronal
        init_neural_memory(user_id)
        
        # Generar token
        token = generate_jwt_token(user_id)
        if not token:
            return jsonify({'error': msgs['error_token']}), 500
        
        print(f"✅ User registered successfully: {user_id}")
        
        return jsonify({
            'success': True,
            'message': msgs['success'],
            'token': token,
            'user_id': user_id,
            'user': {
                'id': user_id,
                'email': email,
                'first_name': data['first_name'],
                'last_name': data['last_name'],
                'subscription_plan': 'free',
                'credits': 100
            },
            'detected_language': detected_language
        }), 201
        
    except Exception as e:
        print(f"❌ Error in register: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/auth/login', methods=['POST'])
def login():
    """Login de usuario - CON DETECCIÓN DE IDIOMA"""
    try:
        print("🚀 Login endpoint called")
        data = request.get_json()
        
        # Detectar idioma del email o cualquier texto proporcionado
        detected_language = detect_user_language(data.get('email', ''))
        
        # Mensajes multiidioma
        messages = {
            'es': {
                'credentials_required': 'Email y contraseña son requeridos',
                'invalid_credentials': 'Credenciales inválidas',
                'use_google': 'Por favor usa Google Sign-In para esta cuenta',
                'error_token': 'Error generando token',
                'success': 'Inicio de sesión exitoso'
            },
            'en': {
                'credentials_required': 'Email and password are required',
                'invalid_credentials': 'Invalid credentials',
                'use_google': 'Please use Google Sign-In for this account',
                'error_token': 'Error generating token',
                'success': 'Login successful'
            }
        }
        
        msgs = messages.get(detected_language, messages['en'])
        
        if not data or 'email' not in data or 'password' not in data:
            return jsonify({'error': msgs['credentials_required']}), 400
        
        email = data['email'].lower().strip()
        password = data['password']
        
        user = get_user_by_email(email)
        if not user:
            return jsonify({'error': msgs['invalid_credentials']}), 401
        
        # Verificar que sea usuario manual (no Google)
        if user.get('auth_provider') != 'manual':
            return jsonify({'error': msgs['use_google']}), 401
        
        if not verify_password(password, user.get('password_hash')):
            return jsonify({'error': msgs['invalid_credentials']}), 401
        
        token = generate_jwt_token(user['id'])
        if not token:
            return jsonify({'error': msgs['error_token']}), 500
        
        print(f"✅ User logged in successfully: {user['id']}")
        
        return jsonify({
            'success': True,
            'message': msgs['success'],
            'token': token,
            'user': {
                'id': user['id'],
                'email': user['email'],
                'first_name': user['first_name'],
                'last_name': user['last_name'],
                'subscription_plan': user['subscription_plan'],
                'credits': user['credits'],
                'auth_provider': user['auth_provider']
            },
            'detected_language': detected_language
        })
        
    except Exception as e:
        print(f"❌ Error in login: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/auth/google', methods=['POST'])
def google_auth():
    """Google auth endpoint"""
    try:
        print("🚀 Google auth endpoint called")
        data = request.get_json()
        
        if not data or not data.get('token'):
            print("❌ No token provided")
            return jsonify({'error': 'Google token is required'}), 400
        
        access_token = data['token']
        print(f"🔑 Received access token (preview): {access_token[:20]}...")
        
        # Verificar token con Google
        user_info = verify_google_access_token(access_token)
        
        if not user_info:
            print("❌ Failed to verify Google token")
            return jsonify({'error': 'Invalid Google token'}), 401
        
        print(f"✅ Google token verified for user: {user_info['email']}")
        
        # Extraer información del usuario
        google_id = user_info['google_id']
        email = user_info['email']
        first_name = user_info['first_name']
        last_name = user_info['last_name']
        
        # Buscar usuario existente por Google ID
        user = get_user_by_google_id(google_id)
        
        if not user:
            # Buscar por email (puede ser un usuario manual existente)
            user = get_user_by_email(email)
            
            if user:
                # Usuario existe pero no tiene Google ID - actualizar
                try:
                    with engine.connect() as conn:
                        conn.execute(
                            text("""
                                UPDATE users 
                                SET google_id = :google_id, auth_provider = 'google', updated_at = NOW()
                                WHERE id = :user_id
                            """),
                            {"google_id": google_id, "user_id": user['id']}
                        )
                        conn.commit()
                        user['google_id'] = google_id
                        user['auth_provider'] = 'google'
                        print(f"✅ Updated existing user with Google ID: {user['id']}")
                except Exception as e:
                    print(f"❌ Error updating user with Google ID: {e}")
            else:
                # Crear nuevo usuario con Google
                user_id = create_user(
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    auth_provider='google',
                    google_id=google_id
                )
                if not user_id:
                    return jsonify({'error': 'Failed to create user'}), 500
                
                # Inicializar memoria neuronal
                init_neural_memory(user_id)
                
                user = get_user_by_id(user_id)
                print(f"✅ Created new Google user: {user_id}")
        
        # Generar token JWT
        token = generate_jwt_token(user['id'])
        
        print(f"🎉 Google authentication successful for: {user['email']}")
        
        return jsonify({
            'success': True,
            'message': 'Google authentication successful',
            'token': token,
            'user': {
                'id': str(user['id']),
                'email': user['email'],
                'first_name': user['first_name'],
                'last_name': user['last_name'],
                'subscription_plan': user['subscription_plan'],
                'credits': user['credits'],
                'auth_provider': user['auth_provider']
            }
        })
        
    except Exception as e:
        print(f"❌ Error in Google auth: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': 'Google authentication failed'}), 500

@app.route('/user/profile', methods=['GET'])
@require_auth
def get_profile(user):
    """Obtiene perfil del usuario"""
    try:
        return jsonify({
            'success': True,
            'user': {
                'id': user['id'],
                'email': user['email'],
                'first_name': user['first_name'],
                'last_name': user['last_name'],
                'subscription_plan': user['subscription_plan'],
                'credits': user['credits'],
                'auth_provider': user['auth_provider'],
                'created_at': user['created_at'].isoformat() if user['created_at'] else None
            }
        })
    except Exception as e:
        print(f"Error getting profile: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/credits/balance', methods=['GET'])
@require_auth
def get_credits_balance(user):
    """Obtiene balance de créditos"""
    try:
        return jsonify({
            'success': True,
            'credits': user['credits'],
            'plan': user['subscription_plan']
        })
    except Exception as e:
        print(f"Error getting credits: {e}")
        return jsonify({'error': str(e)}), 500

# ==============================================================================
#           CHAT ENDPOINTS WITH SESSION MANAGEMENT
# ==============================================================================

@app.route('/chat/new', methods=['POST'])
@require_auth
def create_new_chat(user):
    """
    Crea una nueva conversación asociada a un proyecto
    VERSIÓN ROBUSTA - Maneja JSON malformado del frontend
    """
    try:
        # MANEJO ROBUSTO DEL JSON
        try:
            data = request.get_json(force=True) or {}
        except Exception as json_error:
            print(f"⚠️ JSON parse error: {json_error}")
            print(f"📝 Raw request data: {request.data}")
            print(f"📋 Content-Type: {request.content_type}")
            
            # Si no hay JSON válido, usar valores por defecto
            data = {}
        
        project_id = data.get('project_id')
        
        print(f"🔍 Creating chat for user: {user['id']}")
        print(f"📝 Project ID provided: {project_id}")
        print(f"📊 Request data: {data}")
        
        # Si no hay project_id, buscar o crear uno
        if not project_id:
            print("🔍 No project_id provided, looking for user's projects...")
            
            with engine.connect() as conn:
                # Buscar proyectos existentes del usuario
                project_result = conn.execute(
                    text("""
                        SELECT id FROM projects 
                        WHERE user_id = :user_id 
                        ORDER BY created_at DESC 
                        LIMIT 1
                    """),
                    {"user_id": user['id']}
                ).fetchone()
                
                if project_result:
                    project_id = str(project_result[0])
                    print(f"✅ Using existing project: {project_id}")
                else:
                    # Crear proyecto por defecto
                    project_id = str(uuid.uuid4())
                    print(f"📝 Creating default project: {project_id}")
                    
                    # ✅ INSERT CORREGIDO CON updated_at
                    conn.execute(
                        text("""
                            INSERT INTO projects (
                                id, user_id, project_name, project_description, 
                                kpi_data, status, created_at, updated_at
                            ) VALUES (
                                :id, :user_id, :project_name, :project_description, 
                                :kpi_data, :status, NOW(), NOW()
                            )
                        """),
                        {
                            "id": project_id,
                            "user_id": user['id'],
                            "project_name": "Mi Proyecto Principal",
                            "project_description": "Proyecto creado automáticamente para el chat",
                            "kpi_data": json.dumps({
                                "created_automatically": True,
                                "creation_source": "chat_new_endpoint",
                                "created_at": datetime.now().isoformat()
                            }),
                            "status": "ONBOARDING"
                        }
                    )
                    conn.commit()
                    print(f"✅ Default project created successfully: {project_id}")
        else:
            # Verificar que el proyecto existe y pertenece al usuario
            with engine.connect() as conn:
                project_check = conn.execute(
                    text("""
                        SELECT id, project_name FROM projects 
                        WHERE id = :project_id AND user_id = :user_id
                    """),
                    {"project_id": project_id, "user_id": user['id']}
                ).fetchone()
                
                if not project_check:
                    return jsonify({
                        'error': 'Project not found or access denied',
                        'provided_project_id': project_id,
                        'user_id': user['id'],
                        'suggestion': 'Create a new chat without project_id to auto-create a project'
                    }), 404
                else:
                    print(f"✅ Project verified: {project_check[1]} ({project_id})")
        
        # Generar session_id único para el chat
        session_id = str(uuid.uuid4())
        
        print(f"🎉 New chat session created successfully:")
        print(f"   - User: {user['id']}")
        print(f"   - Project: {project_id}")
        print(f"   - Session: {session_id}")
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'project_id': project_id,
            'message': 'New chat session created successfully',
            'user_id': user['id'],
            'timestamp': datetime.now().isoformat(),
            'debug_info': {
                'received_data': data,
                'content_type': request.content_type,
                'has_json': bool(data)
            }
        })
        
    except Exception as e:
        print(f"❌ Error creating new chat: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'error': 'Could not create new chat session',
            'details': str(e),
            'timestamp': datetime.now().isoformat(),
            'debug_info': {
                'raw_data': str(request.data),
                'content_type': request.content_type,
                'method': request.method
            }
        }), 500

@app.route('/chat/bot', methods=['POST'])
@require_auth
def chat_with_bot(user):
    """
    Procesa mensaje del usuario y retorna respuesta del bot
    VERSIÓN CORREGIDA Y COMPLETA CON DETECCIÓN DE IDIOMA
    """
    try:
        data = request.get_json()
        
        # VALIDACIÓN BÁSICA
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        if not data.get('message'):
            return jsonify({'error': 'Message is required'}), 400
        
        # PROJECT_ID Y SESSION_ID son requeridos
        project_id = data.get('project_id')
        session_id = data.get('session_id')
        
        if not project_id:
            return jsonify({
                'error': 'project_id is required',
                'suggestion': 'Use /chat/new to create a new chat session first'
            }), 400
        
        if not session_id:
            return jsonify({
                'error': 'session_id is required', 
                'suggestion': 'Use /chat/new to create a new chat session first'
            }), 400
        
        # VERIFICAR que el proyecto pertenece al usuario
        with engine.connect() as conn:
            project_result = conn.execute(
                text("""
                    SELECT id, project_name, project_description, kpi_data 
                    FROM projects 
                    WHERE id = :project_id AND user_id = :user_id
                """),
                {"project_id": project_id, "user_id": user['id']}
            ).fetchone()
            
            if not project_result:
                return jsonify({
                    'error': 'Project not found or access denied',
                    'project_id': project_id,
                    'user_id': user['id']
                }), 404
        
        # OBTENER CONTEXTO MEJORADO (incluye detección de idioma)
        enhanced_context = get_enhanced_context_for_chat(user, session_id, project_id, data)
        
        # DETECTAR INTENCIÓN DE BÚSQUEDA DE INVERSORES
        user_message = data.get('message', '')
        detected_language = enhanced_context.get('user_language', 'en')
        
        try:
            wants_investor_search = detect_investor_search_intent(
                user_message, 
                detected_language
            )
        except Exception as search_error:
            print(f"❌ Error detecting investor search intent: {search_error}")
            wants_investor_search = False

        # Detectar si quiere generar documento
    if not wants_investor_search:
        try:
            wants_document = detect_document_generation_intent(user_message, detected_language)
        except Exception as doc_error:
            print(f"❌ Error detecting document intent: {doc_error}")
            wants_document = False
    else:
        wants_document = False  # Si busca inversores, NUNCA puede ser documento

    print(f"📊 Detección final - Inversores: {wants_investor_search}, Documento: {wants_document}")ocument intent: {doc_error}")
            wants_document = False

        # ============================================================================
        #   SI DETECTA GENERACIÓN DE DOCUMENTO - MANEJAR POR SEPARADO
        # ============================================================================
        if wants_document:
            print(f"📄 Usuario quiere generar documento!")
            
            # Identificar tipo de documento
            document_type = identify_document_type(user_message, detected_language)
            
            if not document_type:
                # Preguntar qué documento quiere
                ask_messages = {
                    'es': """¡Perfecto! ¿Qué documento necesitas?

📊 **Pitch Deck** - Presentación para inversores
📋 **Plan de Negocios** - Estrategia completa 
📈 **Plan de Marketing** - Estrategia de crecimiento
💰 **Modelo Financiero** - Proyecciones y números

Solo dime cuál prefieres y lo generaré profesionalmente.""",
                    
                    'en': """Perfect! Which document do you need?

📊 **Pitch Deck** - Presentation for investors
📋 **Business Plan** - Complete strategy
📈 **Marketing Plan** - Growth strategy  
💰 **Financial Model** - Projections and numbers

Just tell me which one and I'll generate it professionally."""
                }
                
                return jsonify({
                    'success': True,
                    'bot': 'interactive_mentor',
                    'response': ask_messages.get(detected_language, ask_messages['en']),
                    'credits_charged': 0,
                    'credits_remaining': get_user_credits(user['id']),
                    'session_id': session_id,
                    'project_id': project_id,
                    'detected_language': detected_language,
                    'awaiting_document_choice': True
                })
            
            # Generar documento profesional
            doc_response = generate_professional_document_in_chat(
                user_id=user['id'],
                project_id=project_id,
                document_type=document_type,
                user_context=enhanced_context,
                user_language=detected_language
            )
            
            # Añadir IDs de sesión y proyecto
            doc_response['session_id'] = session_id
            doc_response['project_id'] = project_id
            doc_response['detected_language'] = detected_language
            
            return jsonify(doc_response)

        # ============================================================================
        #   SI DETECTA BÚSQUEDA DE INVERSORES - MANEJAR POR SEPARADO
        # ============================================================================
        if wants_investor_search:
            print(f"🎯 Usuario quiere buscar inversores!")
            
            # Verificar si tiene plan adecuado
            if user.get('plan', 'free') == 'free':
                # Mensaje de upgrade
                upgrade_messages = {
                    'es': """¡Perfecto! Quieres buscar inversores específicos. 

**Para buscar inversores necesitas el plan Growth:**

🔍 **Con Growth Plan obtienes:**
- Búsqueda ML de 20+ inversores específicos
- Filtros por industria, etapa, ubicación
- 100,000 créditos de lanzamiento
- Solo €20/mes

¿Quieres que upgrade tu plan para buscar inversores ahora?""",
                    'en': """Perfect! You want to search for specific investors.

**To search investors you need Growth Plan:**

🔍 **With Growth Plan you get:**
- ML search of 20+ specific investors  
- Filters by industry, stage, location
- 100,000 launch credits
- Only €20/month

Want to upgrade your plan to search investors now?"""
                }
                
                return jsonify({
                    'success': True,
                    'bot': 'interactive_mentor',
                    'response': upgrade_messages.get(detected_language, upgrade_messages['en']),
                    'credits_charged': 0,
                    'credits_remaining': get_user_credits(user['id']),
                    'session_id': session_id,
                    'project_id': project_id,
                    'detected_language': detected_language,
                    'investor_search_detected': True,
                    'upgrade_required': True
                })
                
            else:
                # ¡TIENE PLAN GROWTH/PRO - EJECUTAR BÚSQUEDA AUTOMÁTICAMENTE!
                print(f"🚀 Ejecutando búsqueda automática de inversores...")
                
                try:
                    # 1. EJECUTAR BÚSQUEDA ML (ya devuelve los 20 mejores por score)
                    search_results = ml_investor_search(
                        query=user_message,
                        user_preferences={
                            'user_id': user['id'],
                            'project_id': project_id,
                            'user_plan': user.get('plan')
                        },
                        max_results=20
                    )
                    
                    if 'error' in search_results:
                        return jsonify({
                            'success': False,
                            'error': 'Error en búsqueda de inversores',
                            'details': search_results['error']
                        }), 500
                    
                    investors_found = search_results.get('results', [])
                    
                    if not investors_found:
                        no_results_messages = {
                            'es': f"""🔍 Busqué inversores con: "{user_message}"

❌ **No encontré inversores** que coincidan exactamente con tu búsqueda.

💡 **Intenta con:**
- Términos más generales: "fintech seed" en lugar de "fintech seed Madrid específico"
- Diferentes etapas: "seed" "series A" "angel"
- Otras ubicaciones: "Europe" "USA" "Spain"

¿Quieres que pruebe con una búsqueda más amplia?""",
                            'en': f"""🔍 Searched for investors with: "{user_message}"

❌ **No investors found** matching your exact search.

💡 **Try with:**
- More general terms: "fintech seed" instead of "fintech seed Madrid specific"
- Different stages: "seed" "series A" "angel"
- Other locations: "Europe" "USA" "Spain"

Want me to try a broader search?"""
                        }
                        
                        return jsonify({
                            'success': True,
                            'bot': 'interactive_mentor',
                            'response': no_results_messages.get(detected_language, no_results_messages['en']),
                            'credits_charged': 0,
                            'credits_remaining': get_user_credits(user['id']),
                            'session_id': session_id,
                            'project_id': project_id,
                            'detected_language': detected_language,
                            'investor_search_executed': True,
                            'investors_found': 0
                        })
                    
                    # 2. CALCULAR CRÉDITOS (10 por resultado - máximo 20 resultados)
                    credits_cost = len(investors_found) * 10
                    user_credits_before = get_user_credits(user['id'])
                    
                    if user_credits_before < credits_cost:
                        return jsonify({
                            'error': 'Créditos insuficientes',
                            'required': credits_cost,
                            'available': user_credits_before
                        }), 402
                    
                    # 3. COBRAR CRÉDITOS
                    credits_after = charge_credits(user['id'], credits_cost)
                    if credits_after is None:
                        return jsonify({'error': 'No se pudieron cobrar créditos'}), 500
                    
                    # 4. GUARDAR INVERSORES EN EL PROYECTO
                    saved_count = save_investors_to_project_simple(
                        user['id'], 
                        project_id, 
                        investors_found
                    )
                    
                    # 5. GUARDAR EN HISTORIAL DE BÚSQUEDAS
                    save_search_to_history_simple(
                        user['id'], 
                        user_message, 
                        len(investors_found), 
                        credits_cost
                    )
                    
                    # 6. CREAR TABLA PARA MOSTRAR EN EL CHAT (solo primeros 10)
                    investors_table = create_investors_table_for_chat(
                        investors_found[:21],  # Solo primeros 10 en chat
                        detected_language
                    )
                    
            
                    
                    return jsonify({
                        'success': True,
                        'bot': 'interactive_mentor',
                        'response': success_messages.get(detected_language, success_messages['en']),
                        'credits_charged': credits_cost,
                        'credits_remaining': credits_after,
                        'session_id': session_id,
                        'project_id': project_id,
                        'detected_language': detected_language,
                        'investor_search_executed': True,
                        'investors_found': len(investors_found),
                        'investors_saved': saved_count,
                        'search_query': user_message,
                        'investors_data': investors_found, 
                        'investors_table': investors_table
                    })
                    
                except Exception as search_error:
                    print(f"❌ Error ejecutando búsqueda: {search_error}")
                    
                    error_messages = {
                        'es': f"""❌ **Error buscando inversores**

Hubo un problema técnico ejecutando la búsqueda.

💡 **Puedes intentar:**
- Reformular tu búsqueda con términos más simples
- Ir manualmente a la página "Buscar Inversores"
- Contactar soporte si el problema persiste""",
                        'en': f"""❌ **Error searching investors**

There was a technical problem executing the search.

💡 **You can try:**
- Rephrase your search with simpler terms
- Go manually to "Search Investors" page
- Contact support if problem persists"""
                    }
                    
                    return jsonify({
                        'success': True,
                        'bot': 'interactive_mentor', 
                        'response': error_messages.get(detected_language, error_messages['en']),
                        'credits_charged': 0,
                        'credits_remaining': get_user_credits(user['id']),
                        'session_id': session_id,
                        'project_id': project_id,
                        'detected_language': detected_language,
                        'investor_search_executed': False,
                        'error_occurred': True
                    })

        # ============================================================================
        #   FLUJO NORMAL DEL BOT (cuando NO quiere buscar inversores ni documentos)
        # ============================================================================
        
        # VERIFICAR CRÉDITOS
        user_credits_before = get_user_credits(user['id'])
        credits_required = CREDIT_COSTS.get('basic_bot', 5)
        
        # Mensajes de error multiidioma
        error_messages = {
            'es': {
                'insufficient_credits': 'Créditos insuficientes',
                'upgrade_needed': 'Necesitas actualizar tu plan'
            },
            'en': {
                'insufficient_credits': 'Insufficient credits',
                'upgrade_needed': 'Upgrade needed'
            }
        }
        
        msgs = error_messages.get(detected_language, error_messages['en'])
        
        if user_credits_before < credits_required:
            return jsonify({
                'error': msgs['insufficient_credits'],
                'required': credits_required,
                'available': user_credits_before,
                'upgrade_needed': True,
                'detected_language': detected_language
            }), 402

        # INYECTAR ROADMAP INTELIGENTE EN CONTEXTO
        if enhanced_context.get('business_context', {}).get('stage'):
            roadmap = bot_manager.get_startup_roadmap(enhanced_context)
            
            # Solo mostrar roadmap cada 5 mensajes o si preguntan por "next steps"
            message_count = enhanced_context.get('total_interactions', 0)
            asks_for_guidance = any(word in user_message.lower() for word in 
                ['next', 'siguiente', 'paso', 'step', 'ahora', 'now', 'qué hago', 'what do'])
            
            if message_count % 5 == 0 or asks_for_guidance:
                enhanced_context['show_roadmap'] = True
                enhanced_context['roadmap'] = roadmap
                         
        # PROCESAR CON BOT MANAGER
        response = bot_manager.process_user_request(
            user_input=user_message,
            user_context=enhanced_context,
            user_id=user['id']
        )
        
        # VERIFICAR RESPUESTA
        if 'error' in response:
            return jsonify(response), 400
        
        # ACTUALIZAR MEMORIA DEL PROYECTO
        if project_id:
            extract_and_update_project_memory(
                user['id'], 
                project_id,
                user_message,
                response.get('response', '')
            )
        
        # PREPARAR RESPUESTA FINAL
        final_response = {
            'success': True,
            'bot': response.get('bot', 'interactive_mentor'),
            'response': response.get('response', ''),
            'credits_charged': response.get('credits_charged_by_bot', credits_required),
            'credits_remaining': get_user_credits(user['id']),
            'session_id': session_id,
            'project_id': project_id,
            'detected_language': detected_language,
            'investor_search_detected': False,  # Siempre False en este flujo
            'processing_success': response.get('processing_success', True)
        }
        
        return jsonify(final_response)
        
    except Exception as e:
        print(f"❌ Error in chat endpoint: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'error': 'Could not process chat',
            'details': str(e)
        }), 500

@app.route('/chat/history', methods=['GET'])
@require_auth
def get_chat_history(user):
    """Obtiene historial de conversaciones agrupadas por sesión"""
    try:
        with engine.connect() as conn:
            # Obtener conversaciones agrupadas por session_id
            result = conn.execute(text("""
                SELECT 
                    COALESCE(session_id, id::text) as conversation_id,
                    MIN(created_at) as started_at,
                    MAX(created_at) as last_message_at,
                    COUNT(*) as message_count,
                    MAX(user_input) as last_message,
                    MAX(bot_used) as last_bot
                FROM neural_interactions 
                WHERE user_id = :user_id 
                GROUP BY COALESCE(session_id, id::text)
                ORDER BY MAX(created_at) DESC
                LIMIT 20
            """), {"user_id": user['id']}).fetchall()
            
            conversations = []
            for row in result:
                conversations.append({
                    'conversation_id': row[0],
                    'started_at': row[1].isoformat(),
                    'last_message_at': row[2].isoformat(),
                    'message_count': row[3],
                    'last_message': row[4][:100] + '...' if len(row[4]) > 100 else row[4],
                    'last_bot': row[5]
                })
        
        return jsonify({
            'success': True,
            'conversations': conversations
        })
        
    except Exception as e:
        print(f"Error getting chat history: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/chat/conversation/<conversation_id>', methods=['GET'])
@require_auth
def get_conversation_messages(user, conversation_id):
    """Obtiene todos los mensajes de una conversación específica"""
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT id, bot_used, user_input, bot_output, credits_charged, 
                       context_data, created_at
                FROM neural_interactions
                WHERE user_id = :user_id 
                AND (session_id = :conversation_id OR id::text = :conversation_id)
                ORDER BY created_at ASC
            """), {
                "user_id": user['id'],
                "conversation_id": conversation_id
            }).fetchall()
            
            messages = []
            for row in result:
                messages.append({
                    'id': str(row[0]),
                    'bot_used': row[1],
                    'user_input': row[2],
                    'bot_output': row[3],
                    'credits_charged': row[4],
                    'context_data': json.loads(row[5]) if row[5] else {},
                    'created_at': row[6].isoformat()
                })
        
        return jsonify({
            'success': True,
            'conversation_id': conversation_id,
            'messages': messages
        })
        
    except Exception as e:
        print(f"Error getting conversation: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/chat/stats', methods=['GET'])
@require_auth
def get_chat_stats(user):
    """Get user's chat statistics"""
    try:
        with engine.connect() as conn:
            # Total interactions
            total_result = conn.execute(text("""
                SELECT COUNT(*) FROM neural_interactions WHERE user_id = :user_id
            """), {"user_id": user['id']}).scalar()
            
            # Bot usage stats
            bot_stats_result = conn.execute(text("""
                SELECT bot_used, COUNT(*) as usage_count, SUM(credits_charged) as total_credits
                FROM neural_interactions 
                WHERE user_id = :user_id 
                GROUP BY bot_used 
                ORDER BY usage_count DESC
            """), {"user_id": user['id']}).fetchall()
            
            # Recent activity (last 7 days)
            recent_activity = conn.execute(text("""
                SELECT DATE(created_at) as date, COUNT(*) as count
                FROM neural_interactions 
                WHERE user_id = :user_id AND created_at > NOW() - INTERVAL '7 days'
                GROUP BY DATE(created_at)
                ORDER BY date DESC
            """), {"user_id": user['id']}).fetchall()
            
            bot_stats = []
            for row in bot_stats_result:
                bot_stats.append({
                    'bot_name': row[0],
                    'usage_count': row[1],
                    'total_credits': row[2]
                })
            
            activity_data = []
            for row in recent_activity:
                activity_data.append({
                    'date': row[0].isoformat(),
                    'count': row[1]
                })
        
        return jsonify({
            'success': True,
            'total_interactions': total_result or 0,
            'bot_usage_stats': bot_stats,
            'recent_activity': activity_data
        })
        
    except Exception as e:
        print(f"Error getting chat stats: {e}")
        return jsonify({'error': 'Failed to get chat statistics'}), 500

# ==============================================================================
#           PROJECT CONTEXT ENDPOINTS
# ==============================================================================

@app.route('/bots/available', methods=['GET'])
@require_auth
def get_available_bots(user):
    """Obtiene lista de bots disponibles"""
    try:
        plan = user['subscription_plan']
        available_bots = []
        
        # Bots básicos (todos los planes)
        available_bots.extend([
            'basic_bot',
            'advanced_bot',
            'expert_bot',
            'document_generation'
        ])
        
        # Bots de Growth plan
        if plan in ['growth', 'pro']:
            available_bots.extend([
                'investor_search_result',
                'employee_search_result'
            ])
        
        # Bots de Pro plan
        if plan == 'pro':
            available_bots.extend([
                'template_generation',
                'unipile_message',
                'automated_sequence'
            ])
        
        return jsonify({
            'success': True,
            'available_bots': available_bots,
            'credit_costs': {bot: CREDIT_COSTS[bot] for bot in available_bots}
        })
        
    except Exception as e:
        print(f"Error getting bots: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/subscription/upgrade', methods=['POST'])
@require_auth
def upgrade_subscription(user):
    """Actualiza el plan de suscripción - CON MENSAJES MULTIIDIOMA"""
    try:
        data = request.get_json()
        
        # Detectar idioma
        detected_language = detect_user_language(str(data))
        
        messages = {
            'es': {
                'plan_required': 'El plan es requerido',
                'invalid_plan': 'Plan inválido',
                'already_subscribed': 'Ya estás suscrito a este plan',
                'upgrade_failed': 'Error al actualizar suscripción',
                'upgrade_success': 'Actualizado exitosamente al plan {}',
                'credits_added': 'créditos añadidos'
            },
            'en': {
                'plan_required': 'Plan is required',
                'invalid_plan': 'Invalid plan',
                'already_subscribed': 'Already subscribed to this plan',
                'upgrade_failed': 'Failed to update subscription',
                'upgrade_success': 'Successfully upgraded to {} plan',
                'credits_added': 'credits added'
            }
        }
        
        msgs = messages.get(detected_language, messages['en'])
        
        if not data or 'plan' not in data:
            return jsonify({'error': msgs['plan_required']}), 400
            
        new_plan = data['plan']
        if new_plan not in SUBSCRIPTION_PLANS:
            return jsonify({'error': msgs['invalid_plan']}), 400
            
        if new_plan == user['subscription_plan']:
            return jsonify({'error': msgs['already_subscribed']}), 400
            
        success = update_subscription_plan(user['id'], new_plan)
        if not success:
            return jsonify({'error': msgs['upgrade_failed']}), 500
            
        return jsonify({
            'success': True,
            'message': msgs['upgrade_success'].format(new_plan),
            'new_plan': new_plan,
            'credits_added': SUBSCRIPTION_PLANS[new_plan]['launch_credits'],
            'detected_language': detected_language
        })
        
    except Exception as e:
        print(f"Error upgrading subscription: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/search/investors', methods=['POST'])
@require_auth
@require_plan('growth')
def search_investors(user):
    """
    Endpoint de búsqueda de inversores - SIEMPRE 20 RESULTADOS MÁXIMO
    """
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        
        # Validar query
        if not query or len(query) < 3:
            return jsonify({
                'error': 'Query demasiado corto (mínimo 3 caracteres)'
            }), 400
        
        print(f"🔍 Búsqueda de {user['email']}: '{query}'")
        
        # CRÉDITOS PARA EXACTAMENTE 20 RESULTADOS
        creditos_necesarios = 20 * CREDIT_COSTS['investor_search_result']
        creditos_usuario = get_user_credits(user['id'])
        
        if creditos_usuario < creditos_necesarios:
            return jsonify({
                'error': 'insufficient_credits',
                'required': creditos_necesarios,
                'available': creditos_usuario,
                'message': f'Necesitas {creditos_necesarios} créditos para buscar inversores'
            }), 402
        
        # Preparar contexto para la búsqueda
        contexto_usuario = {
            'user_id': user['id'],
            'user_plan': user['plan'],
            'preferences': data.get('preferences', {})
        }
        
        # EJECUTAR BÚSQUEDA (SIEMPRE MÁXIMO 20)
        import time
        inicio = time.time()
        
        resultados = ml_investor_search(query, contexto_usuario, 20)
        
        tiempo_procesamiento = round((time.time() - inicio) * 1000)  # en milisegundos
        
        # Verificar si hubo error
        if 'error' in resultados:
            return jsonify(resultados), 500
        
        # COBRAR CRÉDITOS por resultados encontrados
        resultados_encontrados = len(resultados.get('results', []))
        creditos_a_cobrar = resultados_encontrados * CREDIT_COSTS['investor_search_result']
        
        exito_cobro = charge_credits(user['id'], creditos_a_cobrar)
        
        if not exito_cobro:
            return jsonify({
                'error': 'No se pudieron cobrar los créditos'
            }), 500
        
        # Añadir información de créditos y timing
        resultados['credits_charged'] = creditos_a_cobrar
        resultados['credits_remaining'] = get_user_credits(user['id'])
        resultados['processing_time_ms'] = tiempo_procesamiento
        
        print(f"✅ Búsqueda completada: {resultados_encontrados} resultados, {creditos_a_cobrar} créditos, {tiempo_procesamiento}ms")
        
        return jsonify(resultados)
        
    except Exception as e:
        print(f"❌ Error en endpoint de búsqueda: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'error': 'Error en el servicio de búsqueda',
            'details': str(e)
        }), 500
                                 
# ==============================================================================
#           ADMIN ENDPOINTS
# ==============================================================================

@app.route('/admin/users', methods=['GET'])
@require_auth
def get_all_users(user):
    """Obtiene todos los usuarios (solo admin)"""
    try:
        # Verificar que sea admin
        if not user.get('is_admin'):
            return jsonify({'error': 'Admin access required'}), 403
        
        with engine.connect() as conn:
            result = conn.execute(
                text("""
                    SELECT id, email, first_name, last_name, plan, 
                           credits, is_admin, auth_provider, created_at, updated_at
                    FROM users 
                    ORDER BY created_at DESC
                """)
            ).fetchall()
            
            users = []
            for row in result:
                users.append({
                    'id': str(row[0]),
                    'email': row[1],
                    'first_name': row[2],
                    'last_name': row[3],
                    'subscription_plan': row[4],
                    'credits': row[5],
                    'is_admin': row[6],
                    'auth_provider': row[7],
                    'created_at': row[8].isoformat() if row[8] else None,
                    'updated_at': row[9].isoformat() if row[9] else None
                })
        
        return jsonify({
            'success': True,
            'users': users,
            'total_count': len(users)
        })
        
    except Exception as e:
        print(f"❌ Error getting users: {e}")
        return jsonify({'error': 'Could not get users'}), 500
@app.route('/projects/<project_id>/saved-investors', methods=['GET'])
@require_auth
def get_saved_investors_for_sidebar(user, project_id):
    """
    Obtiene inversores guardados para mostrar en el sidebar de la página Investors
    """
    try:
        # Verificar que el proyecto pertenece al usuario
        with engine.connect() as conn:
            project_check = conn.execute(
                text("SELECT 1 FROM projects WHERE id = :project_id AND user_id = :user_id"),
                {"project_id": project_id, "user_id": user['id']}
            ).fetchone()
            
            if not project_check:
                return jsonify({'error': 'Project not found'}), 404
            
            # JOIN para obtener datos completos de inversores guardados
            result = conn.execute(
                text("""
                    SELECT 
                        psi.investor_id,
                        psi.added_at,
                        i."Company_Name",
                        i."Company_Description", 
                        i."Company_Location",
                        i."Investing_Stage",
                        i."Investment_Categories",
                        i."Company_Linkedin",
                        i."Company_Email",
                        i."Company_Website"
                    FROM project_saved_investors psi
                    JOIN investors i ON psi.investor_id = i.id
                    WHERE psi.project_id = :project_id
                    ORDER BY psi.added_at DESC
                """),
                {"project_id": project_id}
            ).fetchall()
        
        # Formatear resultados para el frontend
        saved_investors = []
        for row in result:
            saved_investors.append({
                'investor_id': str(row[0]),
                'added_at': row[1].isoformat(),
                'company_name': row[2] or '',
                'description': row[3] or '',
                'location': row[4] or '',
                'investing_stages': row[5] or '',
                'investment_categories': row[6] or '',
                'linkedin_url': row[7] or '',
                'email': row[8] or '',
                'website': row[9] or ''
            })
        
        return jsonify({
            'success': True,
            'project_id': project_id,
            'saved_investors': saved_investors,
            'total_count': len(saved_investors)
        })
        
    except Exception as e:
        print(f"❌ Error getting saved investors: {e}")
        return jsonify({'error': 'Could not get saved investors'}), 500

@app.route('/projects/<project_id>/search-history', methods=['GET'])
@require_auth 
def get_search_history_for_project(user, project_id):
    """
    Obtiene historial de búsquedas de inversores para mostrar en el sidebar
    """
    try:
        with engine.connect() as conn:
            result = conn.execute(
                text("""
                    SELECT id, query, results_count, credits_spent, created_at, filters_used
                    FROM search_history 
                    WHERE user_id = :user_id 
                    AND search_type = 'investors'
                    ORDER BY created_at DESC
                    LIMIT 20
                """),
                {"user_id": user['id']}
            ).fetchall()
        
        searches = []
        for row in result:
            searches.append({
                'id': str(row[0]),
                'query': row[1],
                'results_count': row[2],
                'credits_spent': row[3],
                'created_at': row[4].isoformat(),
                'filters_used': json.loads(row[5]) if row[5] else {}
            })
        
        return jsonify({
            'success': True,
            'search_history': searches,
            'total_searches': len(searches)
        })
        
    except Exception as e:
        print(f"❌ Error getting search history: {e}")
        return jsonify({'error': 'Could not get search history'}), 500
                        
# ==============================================================================
#           ERROR HANDLERS
# ==============================================================================

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(405)
def method_not_allowed(error):
    return jsonify({'error': 'Method not allowed'}), 405

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500

# ==============================================================================
#           NUEVOS ENDPOINTS
# ==============================================================================

@app.route('/projects', methods=['GET', 'POST'])
@require_auth
def handle_projects(user):
    """Handle projects endpoint - ESQUEMA CORRECTO"""
    try:
        if request.method == 'GET':
            with engine.connect() as conn:
                result = conn.execute(
                    text("""
                        SELECT id, user_id, project_name, project_description, kpi_data, status, created_at
                        FROM projects 
                        WHERE user_id = :user_id
                        ORDER BY created_at DESC
                    """),
                    {"user_id": user['id']}
                ).fetchall()
                
                projects = []
                for row in result:
                    try:
                        kpi_data = json.loads(row[4]) if row[4] else {}
                    except:
                        kpi_data = {}
                    
                    projects.append({
                        'id': str(row[0]),
                        'user_id': str(row[1]),
                        'project_name': row[2],
                        'project_description': row[3],
                        'kpi_data': kpi_data,
                        'status': row[5],
                        'created_at': row[6].isoformat() if row[6] else None
                    })
            
            return jsonify({
                'success': True,
                'projects': projects,
                'count': len(projects)
            })
        
        elif request.method == 'POST':
            data = request.get_json()
            if not data:
                return jsonify({'error': 'No data provided'}), 400
            
            # Detectar idioma para mensajes
            detected_language = detect_user_language(data.get('name', '') + ' ' + data.get('description', ''))
            
            project_id = str(uuid.uuid4())
            project_name = data.get('name', 'Untitled Project')
            project_description = data.get('description', '')
            
            project_data = {
                'name': data.get('name', ''),
                'description': data.get('description', ''),
                'industry': data.get('industry', ''),
                'stage': data.get('stage', ''),
                'location': data.get('location', ''),
                'website': data.get('website', ''),
                'created_by': user['email']
            }
            
            with engine.connect() as conn:
                # ✅ INSERT CORREGIDO - Incluye updated_at
                conn.execute(
                    text("""
                        INSERT INTO projects (
                            id, user_id, project_name, project_description, kpi_data, status, created_at, updated_at
                        ) VALUES (
                            :id, :user_id, :project_name, :project_description, :kpi_data, :status, NOW(), NOW()
                        )
                    """),
                    {
                        "id": project_id,
                        "user_id": user['id'],
                        "project_name": project_name,
                        "project_description": project_description,
                        "kpi_data": json.dumps(project_data),
                        "status": "ONBOARDING"
                    }
                )
                conn.commit()
                print(f"✅ Project created successfully: {project_id}")
            
            success_messages = {
                'es': 'Proyecto creado exitosamente',
                'en': 'Project created successfully'
            }
            
            return jsonify({
                'success': True,
                'message': success_messages.get(detected_language, success_messages['en']),
                'project_id': project_id,
                'detected_language': detected_language
            })
            
    except Exception as e:
        print(f"❌ Error in projects endpoint: {e}")
        return jsonify({'error': f'Database error: {str(e)}'}), 500

@app.route('/chat/recent', methods=['GET'])
@require_auth
def get_recent_chats_with_titles(user):
    """Obtiene chats recientes con títulos - OPTIMIZADO para tu esquema exacto"""
    try:
        limit = min(int(request.args.get('limit', 20)), 50)
        project_id = request.args.get('project_id')
        
        base_query = """
            SELECT 
                COALESCE(ni.session_id::text, ni.id::text) as session_id,
                ni.project_id,
                p.project_name,
                MIN(ni.created_at) as started_at,
                MAX(ni.created_at) as last_message_at,
                COUNT(*) as message_count,
                (array_agg(ni.user_input ORDER BY ni.created_at ASC))[1] as first_message,
                (array_agg(ni.user_input ORDER BY ni.created_at DESC))[1] as last_message_preview,
                (array_agg(ni.bot_used ORDER BY ni.created_at DESC))[1] as last_bot_used,
                MAX(ni.session_title) as session_title
            FROM neural_interactions ni
            LEFT JOIN projects p ON ni.project_id = p.id
            WHERE ni.user_id = :user_id
        """
        
        params = {"user_id": user['id'], "limit": limit}
        
        if project_id:
            base_query += " AND ni.project_id = :project_id"
            params["project_id"] = project_id
        
        base_query += """
            GROUP BY COALESCE(ni.session_id::text, ni.id::text), ni.project_id, p.project_name
            ORDER BY MAX(ni.created_at) DESC
            LIMIT :limit
        """
        
        with engine.connect() as conn:
            result = conn.execute(text(base_query), params).fetchall()
            
            chats = []
            for row in result:
                # Si no hay título, generarlo ahora
                display_title = row[9] if row[9] else generate_simple_title_from_message(row[6], row[7])
                
                chats.append({
                    'session_id': row[0],
                    'project_id': str(row[1]) if row[1] else None,
                    'project_name': row[2] or 'Sin proyecto',
                    'title': display_title,  # ¡TÍTULO AQUÍ!
                    'started_at': row[3].isoformat(),
                    'last_message_at': row[4].isoformat(),
                    'message_count': row[5],
                    'last_message_preview': (row[7][:60] + '...') if len(row[7] or '') > 60 else (row[7] or ''),
                    'last_bot_used': row[8] or 'unknown'
                })
        
        return jsonify({
            'success': True,
            'chats': chats,
            'total_count': len(chats),
            'filtered_by_project': project_id is not None
        })
        
    except Exception as e:
        print(f"❌ Error getting recent chats: {e}")
        return jsonify({'error': 'Could not get recent chats'}), 500

@app.route('/chat/messages/<session_id>', methods=['GET'])
@require_auth
def get_chat_messages(user, session_id):
    """
    Obtiene todos los mensajes de una sesión/conversación específica
    
    REQUEST: GET /chat/messages/uuid-de-sesion
    """
    try:
        with engine.connect() as conn:
            result = conn.execute(text("""
                SELECT ni.id, ni.bot_used, ni.user_input, ni.bot_output, ni.credits_charged, 
                       ni.context_data, ni.created_at, ni.project_id, p.project_name
                FROM neural_interactions ni
                LEFT JOIN projects p ON ni.project_id = p.id
                WHERE ni.user_id = :user_id 
                AND (ni.session_id::text = :session_id OR ni.id::text = :session_id)
                ORDER BY ni.created_at ASC
            """), {
                "user_id": user['id'],
                "session_id": session_id
            }).fetchall()
            
            if not result:
                return jsonify({
                    'success': True,
                    'session_id': session_id,
                    'messages': [],
                    'total_messages': 0,
                    'message': 'No messages found for this session'
                })
            
            messages = []
            for row in result:
                try:
                    context_data = json.loads(row[5]) if row[5] else {}
                except:
                    context_data = {}
                
                messages.append({
                    'id': str(row[0]),
                    'user_message': row[2],
                    'bot_response': row[3],
                    'bot_used': row[1],
                    'credits_charged': row[4],
                    'timestamp': row[6].isoformat(),
                    'context': context_data
                })
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'messages': messages,
            'total_messages': len(messages)
        })
        
    except Exception as e:
        print(f"❌ Error getting chat messages: {e}")
        return jsonify({'error': 'Could not get chat messages'}), 500

@app.route('/chat/delete/<session_id>', methods=['DELETE'])
@require_auth
def delete_chat_session(user, session_id):
    """
    Elimina una conversación completa - SQL CORREGIDO
    """
    try:
        with engine.connect() as conn:
            # ✅ QUERY CORREGIDO - Sin referencia a tabla 'ni'
            count_result = conn.execute(text("""
                SELECT COUNT(*) FROM neural_interactions
                WHERE user_id = :user_id 
                AND (session_id::text = :session_id OR id::text = :session_id)
            """), {"user_id": user['id'], "session_id": session_id}).scalar()
            
            if count_result == 0:
                # Detectar idioma para mensaje
                detected_language = 'es'  # Default, ya que no hay mensaje del usuario
                error_messages = {
                    'es': 'Sesión de chat no encontrada',
                    'en': 'Chat session not found'
                }
                return jsonify({'error': error_messages.get(detected_language, error_messages['en'])}), 404
            
            # ✅ DELETE CORREGIDO - Sin referencia a tabla 'ni'
            conn.execute(text("""
                DELETE FROM neural_interactions
                WHERE user_id = :user_id 
                AND (session_id::text = :session_id OR id::text = :session_id)
            """), {"user_id": user['id'], "session_id": session_id})
            
            conn.commit()
        
        success_messages = {
            'es': 'Sesión de chat eliminada exitosamente',
            'en': 'Chat session deleted successfully'
        }
        
        return jsonify({
            'success': True,
            'message': success_messages.get('es', success_messages['en']),
            'deleted_messages': count_result
        })
        
    except Exception as e:
        print(f"❌ Error deleting chat session: {e}")
        return jsonify({'error': 'Could not delete chat session'}), 500

@app.route('/documents/generate', methods=['POST'])
@require_auth
def generate_document_with_bot(user):
    """Genera documento extenso usando bot específico - CON DETECCIÓN DE IDIOMA"""
    try:
        data = request.get_json()
        
        # Detectar idioma
        user_requirements = data.get('requirements', '')
        detected_language = detect_user_language(user_requirements)
        
        messages = {
            'es': {
                'required_fields': 'document_type y project_id son requeridos',
                'project_not_found': 'Proyecto no encontrado',
                'invalid_doc_type': 'Tipo de documento inválido',
                'insufficient_credits': 'Créditos insuficientes',
                'error_generating': 'No se pudo generar el documento'
            },
            'en': {
                'required_fields': 'document_type and project_id are required',
                'project_not_found': 'Project not found',
                'invalid_doc_type': 'Invalid document type',
                'insufficient_credits': 'Insufficient credits',
                'error_generating': 'Could not generate document'
            }
        }
        
        msgs = messages.get(detected_language, messages['en'])
        
        if not data or not data.get('document_type') or not data.get('project_id'):
            return jsonify({'error': msgs['required_fields']}), 400
        
        document_type = data['document_type']  # 'pitch_deck', 'business_plan', 'marketing_plan'
        project_id = data['project_id']
        
        # Verificar proyecto
        with engine.connect() as conn:
            project_result = conn.execute(
                text("SELECT * FROM projects WHERE id = :project_id AND user_id = :user_id"),
                {"project_id": project_id, "user_id": user['id']}
            ).fetchone()
            
            if not project_result:
                return jsonify({'error': msgs['project_not_found']}), 404
        
        # Determinar bot y créditos requeridos
        bot_config = {
            'pitch_deck': {'bot': 'pitch_deck_master', 'credits': 100},
            'business_plan': {'bot': 'strategy_consultant', 'credits': 150},
            'marketing_plan': {'bot': 'content_machine', 'credits': 120},
            'financial_model': {'bot': 'financial_modeler', 'credits': 130}
        }
        
        if document_type not in bot_config:
            return jsonify({'error': msgs['invalid_doc_type']}), 400
        
        bot_id = bot_config[document_type]['bot']
        credits_required = bot_config[document_type]['credits']
        
        # Verificar créditos
        if not has_sufficient_credits(user['id'], credits_required):
            return jsonify({
                'error': msgs['insufficient_credits'],
                'required': credits_required,
                'available': get_user_credits(user['id'])
            }), 402
        
        # Crear prompt especializado para documento extenso EN EL IDIOMA DETECTADO
        project_data = json.loads(project_result[4]) if project_result[4] else {}
        
        # Prompts multiidioma para documentos
        doc_prompts = {
            'es': {
                'pitch_deck': f"""
                Eres el mejor creador de pitch decks del mundo. Crea un pitch deck COMPLETO y DETALLADO para esta startup.
                
                INFORMACIÓN DEL PROYECTO:
                - Nombre: {project_data.get('name', 'Mi Startup')}
                - Industria: {project_data.get('industry', 'Tecnología')}
                - Descripción: {project_data.get('description', '')}
                - Etapa: {project_data.get('stage', 'Seed')}
                
                REQUISITOS ESPECÍFICOS:
                {user_requirements}
                
                ESTRUCTURA REQUERIDA (crear cada sección en detalle):
                1. **PORTADA**: Nombre, tagline, logo placeholder
                2. **PROBLEMA**: Problema específico y dolor real
                3. **SOLUCIÓN**: Solución única y diferenciada
                4. **TAMAÑO DE MERCADO**: TAM, SAM, SOM con números reales
                5. **PRODUCTO**: Características clave y demostración
                6. **MODELO DE NEGOCIO**: Cómo generas dinero
                7. **TRACCIÓN**: Métricas y logros actuales
                8. **COMPETENCIA**: Análisis competitivo
                9. **MARKETING**: Estrategia go-to-market
                10. **EQUIPO**: Fundadores y equipo clave
                11. **FINANZAS**: Proyecciones 3-5 años
                12. **FINANCIACIÓN**: Cantidad, uso de fondos, valoración
                13. **CRONOGRAMA**: Roadmap y milestones
                14. **APÉNDICE**: Información adicional
                
                FORMATO:
                - Cada slide con título H2
                - Contenido detallado y específico
                - Números y métricas concretas
                - Call-to-action en cada slide
                - Mínimo 2000 palabras total
                
                Responde COMPLETAMENTE EN ESPAÑOL.
                """,
                
                'business_plan': f"""
                Crea un PLAN DE NEGOCIO COMPLETO y PROFESIONAL para esta startup.
                
                INFORMACIÓN DEL PROYECTO:
                - Nombre: {project_data.get('name', 'Mi Startup')}
                - Industria: {project_data.get('industry', 'Tecnología')}
                - Descripción: {project_data.get('description', '')}
                
                REQUISITOS:
                {user_requirements}
                
                ESTRUCTURA COMPLETA:
                
                ## 1. RESUMEN EJECUTIVO
                - Resumen ejecutivo de 2 páginas
                - Propuesta de valor única
                - Proyecciones financieras clave
                - Financiación necesaria
                
                ## 2. DESCRIPCIÓN DE LA EMPRESA
                - Historia y misión
                - Visión y valores
                - Estructura legal
                - Ubicación y operaciones
                
                ## 3. ANÁLISIS DE MERCADO
                - Análisis de industria
                - Segmentación del mercado objetivo
                - Tamaño de mercado (TAM, SAM, SOM)
                - Tendencias y oportunidades
                
                ## 4. ANÁLISIS COMPETITIVO
                - Panorama competitivo
                - Competidores directos vs indirectos
                - Análisis SWOT
                - Ventaja competitiva sostenible
                
                ## 5. PRODUCTOS Y SERVICIOS
                - Descripción detallada del producto
                - Características y beneficios
                - Roadmap de desarrollo
                - Propiedad intelectual
                
                ## 6. ESTRATEGIA DE MARKETING Y VENTAS
                - Personas de clientes
                - Marketing mix (4Ps)
                - Embudo de ventas
                - Estrategia de adquisición de clientes
                - Estrategia de precios
                
                ## 7. PLAN DE OPERACIONES
                - Flujo de trabajo operacional
                - Cadena de suministro
                - Infraestructura tecnológica
                - Control de calidad
                
                ## 8. EQUIPO DE GESTIÓN
                - Biografías del equipo y experiencia
                - Organigrama
                - Junta asesora
                - Plan de contratación
                
                ## 9. PROYECCIONES FINANCIERAS
                - Proyección P&L a 5 años
                - Análisis de flujo de caja
                - Análisis de punto de equilibrio
                - Ratios financieros clave
                - Requisitos de financiación
                
                ## 10. ANÁLISIS DE RIESGOS
                - Riesgos de mercado
                - Riesgos operacionales
                - Riesgos financieros
                - Estrategias de mitigación
                
                Mínimo 4000 palabras con números específicos y análisis detallado.
                Responde COMPLETAMENTE EN ESPAÑOL.
                """
            },
            'en': {
                'pitch_deck': f"""
                You are the world's best pitch deck creator. Create a COMPLETE and DETAILED pitch deck for this startup.
                
                PROJECT INFORMATION:
                - Name: {project_data.get('name', 'My Startup')}
                - Industry: {project_data.get('industry', 'Technology')}
                - Description: {project_data.get('description', '')}
                - Stage: {project_data.get('stage', 'Seed')}
                
                SPECIFIC REQUIREMENTS:
                {user_requirements}
                
                REQUIRED STRUCTURE (create each section in detail):
                1. **COVER SLIDE**: Name, tagline, logo placeholder
                2. **PROBLEM**: Specific problem and real pain
                3. **SOLUTION**: Unique and differentiated solution
                4. **MARKET SIZE**: TAM, SAM, SOM with real numbers
                5. **PRODUCT**: Key features and demonstration
                6. **BUSINESS MODEL**: How you make money
                7. **TRACTION**: Current metrics and achievements
                8. **COMPETITION**: Competitive analysis
                9. **MARKETING**: Go-to-market strategy
                10. **TEAM**: Founders and key team
                11. **FINANCIALS**: 3-5 year projections
                12. **FUNDING**: Amount, use of funds, valuation
                13. **TIMELINE**: Roadmap and milestones
                14. **APPENDIX**: Additional information
                
                FORMAT:
                - Each slide with H2 title
                - Detailed and specific content
                - Concrete numbers and metrics
                - Call-to-action on each slide
                - Minimum 2000 words total
                
                Respond COMPLETELY IN ENGLISH.
                """,
                
                'business_plan': f"""
                Create a COMPLETE and PROFESSIONAL BUSINESS PLAN for this startup.
                
                PROJECT INFORMATION:
                - Name: {project_data.get('name', 'My Startup')}
                - Industry: {project_data.get('industry', 'Technology')}
                - Description: {project_data.get('description', '')}
                
                REQUIREMENTS:
                {user_requirements}
                
                COMPLETE STRUCTURE:
                
                ## 1. EXECUTIVE SUMMARY
                - 2-page executive summary
                - Unique value proposition
                - Key financial projections
                - Funding needed
                
                ## 2. COMPANY DESCRIPTION
                - History and mission
                - Vision and values
                - Legal structure
                - Location and operations
                
                ## 3. MARKET ANALYSIS
                - Industry analysis
                - Target market segmentation
                - Market size (TAM, SAM, SOM)
                - Trends and opportunities
                
                ## 4. COMPETITIVE ANALYSIS
                - Competitive landscape
                - Direct vs indirect competitors
                - SWOT analysis
                - Sustainable competitive advantage
                
                ## 5. PRODUCTS & SERVICES
                - Detailed product description
                - Features and benefits
                - Development roadmap
                - Intellectual property
                
                ## 6. MARKETING & SALES STRATEGY
                - Customer personas
                - Marketing mix (4Ps)
                - Sales funnel
                - Customer acquisition strategy
                - Pricing strategy
                
                ## 7. OPERATIONS PLAN
                - Operational workflow
                - Supply chain
                - Technology infrastructure
                - Quality control
                
                ## 8. MANAGEMENT TEAM
                - Team bios and experience
                - Organizational chart
                - Advisory board
                - Hiring plan
                
                ## 9. FINANCIAL PROJECTIONS
                - 5-year P&L projection
                - Cash flow analysis
                - Break-even analysis
                - Key financial ratios
                - Funding requirements
                
                ## 10. RISK ANALYSIS
                - Market risks
                - Operational risks
                - Financial risks
                - Mitigation strategies
                
                Minimum 4000 words with specific numbers and detailed analysis.
                Respond COMPLETELY IN ENGLISH.
                """
            }
        }
        
        # Seleccionar prompts según idioma
        prompts = doc_prompts.get(detected_language, doc_prompts['en'])[document_type]
        
        # Generar documento con Gemini
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            prompts,
            generation_config={
                "temperature": 0.7,
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": 8000,  # Máximo para documentos extensos
            }
        )
        
        # Cobrar créditos
        credits_after = charge_credits(user['id'], credits_required)
        if credits_after is None:
            return jsonify({'error': msgs['error_generating']}), 500
        
        # Guardar documento
        document_id = str(uuid.uuid4())
        title = f"{document_type.replace('_', ' ').title()} - {project_data.get('name', 'Mi Startup')}"
        
        with engine.connect() as conn:
            conn.execute(
                text("""
                    INSERT INTO generated_documents (
                        id, user_id, bot_used, document_type, title, content, 
                        format, metadata, credits_used, created_at, updated_at
                    ) VALUES (
                        :id, :user_id, :bot_used, :document_type, :title, :content,
                        'markdown', :metadata, :credits_used, NOW(), NOW()
                    )
                """),
                {
                    "id": document_id,
                    "user_id": user['id'],
                    "bot_used": bot_id,
                    "document_type": document_type,
                    "title": title,
                    "content": response.text,
                    "metadata": json.dumps({
                        "project_id": project_id,
                        "word_count": len(response.text.split()),
                        "generated_at": datetime.now().isoformat(),
                        "user_requirements": user_requirements,
                        "language": detected_language
                    }),
                    "credits_used": credits_required
                }
            )
            conn.commit()
        
        return jsonify({
            'success': True,
            'document_id': document_id,
            'title': title,
            'document_type': document_type,
            'content_preview': response.text[:500] + '...',
            'word_count': len(response.text.split()),
            'credits_used': credits_required,
            'credits_remaining': credits_after,
            'download_url': f'/documents/{document_id}/download',
            'view_url': f'/documents/{document_id}/view',
            'detected_language': detected_language
        })
        
    except Exception as e:
        print(f"❌ Error generating document: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Could not generate document: {str(e)}'}), 500

@app.route('/documents/<document_id>/view', methods=['GET'])
@require_auth
def view_document(user, document_id):
    """Ver contenido completo del documento"""
    try:
        with engine.connect() as conn:
            result = conn.execute(
                text("""
                    SELECT title, content, document_type, created_at, metadata, bot_used
                    FROM generated_documents 
                    WHERE id = :document_id AND user_id = :user_id
                """),
                {"document_id": document_id, "user_id": user['id']}
            ).fetchone()
            
            if not result:
                return jsonify({'error': 'Document not found'}), 404
            
            # Incrementar contador de views
            conn.execute(
                text("UPDATE generated_documents SET download_count = download_count + 1 WHERE id = :document_id"),
                {"document_id": document_id}
            )
            conn.commit()
        
        return jsonify({
            'success': True,
            'document_id': document_id,
            'title': result[0],
            'content': result[1],
            'document_type': result[2],
            'created_at': result[3].isoformat(),
            'metadata': json.loads(result[4]) if result[4] else {},
            'bot_used': result[5]
        })
        
    except Exception as e:
        print(f"❌ Error viewing document: {e}")
        return jsonify({'error': 'Could not view document'}), 500

# Configuración para subida de archivos
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx', 'xls', 'xlsx', 'csv', 'png', 'jpg', 'jpeg'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/files/upload', methods=['POST'])
@require_auth
def upload_file(user):
    """Subir archivo para análisis - CON DETECCIÓN DE IDIOMA"""
    try:
        # Detectar idioma del usuario
        detected_language = 'es'  # Default
        
        messages = {
            'es': {
                'no_file': 'No se encontró archivo',
                'no_filename': 'Archivo sin nombre',
                'invalid_type': 'Tipo de archivo no permitido',
                'file_too_large': 'Archivo demasiado grande (máximo 10MB)',
                'upload_failed': 'Error al subir archivo',
                'upload_success': 'Archivo subido exitosamente'
            },
            'en': {
                'no_file': 'No file found',
                'no_filename': 'File without name',
                'invalid_type': 'File type not allowed',
                'file_too_large': 'File too large (maximum 10MB)',
                'upload_failed': 'Error uploading file',
                'upload_success': 'File uploaded successfully'
            }
        }
        
        msgs = messages.get(detected_language, messages['en'])
        
        if 'file' not in request.files:
            return jsonify({'error': msgs['no_file']}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': msgs['no_filename']}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': msgs['invalid_type']}), 400
        
        # Verificar tamaño
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        
        if file_size > MAX_FILE_SIZE:
            return jsonify({'error': msgs['file_too_large']}), 400
        
        # Procesar archivo
        filename = secure_filename(file.filename)
        file_id = str(uuid.uuid4())
        
        # Extraer contenido según tipo
        content = ""
        file_type = filename.rsplit('.', 1)[1].lower()
        
        if file_type == 'pdf':
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                content += page.extract_text()
        
        elif file_type in ['doc', 'docx']:
            doc = docx.Document(file)
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        elif file_type in ['txt', 'csv']:
            content = file.read().decode('utf-8')
        
        # Guardar metadata en DB
        with engine.connect() as conn:
            conn.execute(
                text("""
                    INSERT INTO uploaded_files (
                        id, user_id, filename, file_type, file_size, 
                        content_preview, metadata, created_at
                    ) VALUES (
                        :id, :user_id, :filename, :file_type, :file_size,
                        :content_preview, :metadata, NOW()
                    )
                """),
                {
                    "id": file_id,
                    "user_id": user['id'],
                    "filename": filename,
                    "file_type": file_type,
                    "file_size": file_size,
                    "content_preview": content[:1000],
                    "metadata": json.dumps({
                        "original_filename": file.filename,
                        "content_length": len(content),
                        "upload_timestamp": datetime.now().isoformat()
                    })
                }
            )
            conn.commit()
        
        # Analizar contenido con Gemini
        analysis = analyze_file_content(content, file_type, detected_language)
        
        return jsonify({
            'success': True,
            'message': msgs['upload_success'],
            'file_id': file_id,
            'filename': filename,
            'file_type': file_type,
            'file_size': file_size,
            'content_preview': content[:500] + '...' if len(content) > 500 else content,
            'analysis': analysis,
            'detected_language': detected_language
        })
        
    except Exception as e:
        print(f"❌ Error uploading file: {e}")
        return jsonify({'error': msgs['upload_failed']}), 500

def analyze_file_content(content, file_type, user_language='en'):
    """Analiza contenido del archivo con Gemini - MULTIIDIOMA"""
    try:
        analysis_prompts = {
            'es': f"""
            Analiza el siguiente contenido de un archivo {file_type} y proporciona:
            1. Resumen ejecutivo (2-3 párrafos)
            2. Puntos clave identificados
            3. Recomendaciones para una startup
            4. Posibles casos de uso
            
            Contenido:
            {content[:3000]}
            
            Responde en español con formato estructurado.
            """,
            
            'en': f"""
            Analyze the following content from a {file_type} file and provide:
            1. Executive summary (2-3 paragraphs)
            2. Key points identified
            3. Recommendations for a startup
            4. Possible use cases
            
            Content:
            {content[:3000]}
            
            Respond in English with structured format.
            """
        }
        
        prompt = analysis_prompts.get(user_language, analysis_prompts['en'])
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.5,
                "max_output_tokens": 1000,
            }
        )
        
        return {
            'summary': response.text,
            'analyzed_at': datetime.now().isoformat(),
            'content_type': file_type,
            'language': user_language
        }
        
    except Exception as e:
        print(f"❌ Error analyzing file: {e}")
        return {
            'summary': 'Could not analyze file content',
            'error': str(e)
        }

@app.route('/feedback', methods=['POST'])
@require_auth
def submit_feedback(user):
    """Enviar feedback del usuario - CON DETECCIÓN DE IDIOMA"""
    try:
        data = request.get_json()
        
        # Detectar idioma del feedback
        feedback_text = data.get('feedback', '')
        detected_language = detect_user_language(feedback_text)
        
        messages = {
            'es': {
                'feedback_required': 'El feedback es requerido',
                'feedback_saved': 'Gracias por tu feedback'
            },
            'en': {
                'feedback_required': 'Feedback is required',
                'feedback_saved': 'Thank you for your feedback'
            }
        }
        
        msgs = messages.get(detected_language, messages['en'])
        
        if not data or not feedback_text:
            return jsonify({'error': msgs['feedback_required']}), 400
        
        # Guardar feedback
        feedback_id = str(uuid.uuid4())
        
        with engine.connect() as conn:
            conn.execute(
                text("""
                    INSERT INTO user_feedback (
                        id, user_id, feedback_type, feedback_text, 
                        rating, metadata, created_at
                    ) VALUES (
                        :id, :user_id, :feedback_type, :feedback_text,
                        :rating, :metadata, NOW()
                    )
                """),
                {
                    "id": feedback_id,
                    "user_id": user['id'],
                    "feedback_type": data.get('type', 'general'),
                    "feedback_text": feedback_text,
                    "rating": data.get('rating'),
                    "metadata": json.dumps({
                        "source": data.get('source', 'web'),
                        "context": data.get('context', {}),
                        "user_agent": request.headers.get('User-Agent', ''),
                        "detected_language": detected_language
                    })
                }
            )
            conn.commit()
        
        # Analizar sentimiento del feedback
        sentiment = analyze_feedback_sentiment(feedback_text, detected_language)
        
        return jsonify({
            'success': True,
            'message': msgs['feedback_saved'],
            'feedback_id': feedback_id,
            'sentiment_analysis': sentiment,
            'detected_language': detected_language
        })
        
    except Exception as e:
        print(f"❌ Error submitting feedback: {e}")
        return jsonify({'error': 'Could not submit feedback'}), 500

def analyze_feedback_sentiment(feedback_text, user_language='en'):
    """Analiza sentimiento del feedback con Gemini - MULTIIDIOMA"""
    try:
        sentiment_prompts = {
            'es': f"""
            Analiza el sentimiento del siguiente feedback:
            
            "{feedback_text}"
            
            Responde con:
            1. Sentimiento general: Positivo/Neutral/Negativo
            2. Puntuación de sentimiento: 0-10
            3. Temas principales mencionados
            4. Sugerencias de mejora identificadas
            
            Formato JSON en español.
            """,
            
            'en': f"""
            Analyze the sentiment of the following feedback:
            
            "{feedback_text}"
            
            Respond with:
            1. Overall sentiment: Positive/Neutral/Negative
            2. Sentiment score: 0-10
            3. Main topics mentioned
            4. Improvement suggestions identified
            
            JSON format in English.
            """
        }
        
        prompt = sentiment_prompts.get(user_language, sentiment_prompts['en'])
        
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.3,
                "max_output_tokens": 500,
            }
        )
        
        # Intentar parsear como JSON
        try:
            response_text = response.text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:-3]
            elif response_text.startswith('```'):
                response_text = response_text[3:-3]
            
            return json.loads(response_text)
        except:
            return {
                'sentiment': 'analyzed',
                'raw_analysis': response.text
            }
        
    except Exception as e:
        print(f"❌ Error analyzing sentiment: {e}")
        return {
            'sentiment': 'error',
            'error': str(e)
        }

@app.route('/analytics/usage', methods=['GET'])
@require_auth
def get_usage_analytics(user):
    """Obtener analytics de uso del usuario"""
    try:
        with engine.connect() as conn:
            # Uso por día (últimos 30 días)
            daily_usage = conn.execute(
                text("""
                    SELECT DATE(created_at) as date, 
                           COUNT(*) as interactions,
                           SUM(credits_charged) as credits_used
                    FROM neural_interactions
                    WHERE user_id = :user_id 
                    AND created_at > NOW() - INTERVAL '30 days'
                    GROUP BY DATE(created_at)
                    ORDER BY date DESC
                """),
                {"user_id": user['id']}
            ).fetchall()
            
            # Uso por bot
            bot_usage = conn.execute(
                text("""
                    SELECT bot_used, 
                           COUNT(*) as times_used,
                           SUM(credits_charged) as total_credits
                    FROM neural_interactions
                    WHERE user_id = :user_id
                    GROUP BY bot_used
                    ORDER BY times_used DESC
                """),
                {"user_id": user['id']}
            ).fetchall()
            
            # Documentos generados
            documents = conn.execute(
                text("""
                    SELECT document_type, 
                           COUNT(*) as count,
                           SUM(credits_used) as credits_spent
                    FROM generated_documents
                    WHERE user_id = :user_id
                    GROUP BY document_type
                """),
                {"user_id": user['id']}
            ).fetchall()
        
        # Formatear datos
        daily_data = [{
            'date': row[0].isoformat(),
            'interactions': row[1],
            'credits_used': row[2] or 0
        } for row in daily_usage]
        
        bot_data = [{
            'bot': row[0],
            'times_used': row[1],
            'total_credits': row[2] or 0
        } for row in bot_usage]
        
        document_data = [{
            'type': row[0],
            'count': row[1],
            'credits_spent': row[2] or 0
        } for row in documents]
        
        return jsonify({
            'success': True,
            'analytics': {
                'daily_usage': daily_data,
                'bot_usage': bot_data,
                'documents_generated': document_data,
                'total_credits_used': sum(b['total_credits'] for b in bot_data),
                'total_interactions': sum(d['interactions'] for d in daily_data),
                'current_credits': user['credits']
            }
        })
        
    except Exception as e:
        print(f"❌ Error getting analytics: {e}")
        return jsonify({'error': 'Could not get analytics'}), 500

def ml_investor_search(query, user_preferences, max_results=21):
    """
    Función principal de búsqueda ML de inversores - VERSIÓN CON 21 RESULTADOS
    """
    try:
        print(f"🔍 Iniciando búsqueda ML: '{query}'")
        
        # Verificar que engine esté disponible
        if not engine:
            print("❌ Engine de base de datos no disponible")
            return {
                "error": "Database connection not available",
                "success": False,
                "query": query
            }
        
        # Crear instancia del motor de búsqueda
        search_engine = InvestorSearchSimple(engine)
        
        # Ejecutar búsqueda (ahora devuelve 21 con scores 72-99)
        resultado = search_engine.buscar_inversores(query)
        
        print(f"✅ Búsqueda completada: {len(resultado.get('results', []))} resultados")
        return resultado
        
    except Exception as e:
        print(f"❌ Error en ml_investor_search: {e}")
        import traceback
        traceback.print_exc()
        return {
            "error": f"Search failed: {str(e)}",
            "success": False,
            "query": query,
            "details": str(e)
        }

def require_plan(required_plan):
    """
    Decorator para verificar que el usuario tiene el plan requerido
    """
    def decorator(f):
        @wraps(f)
        def decorated_function(user, *args, **kwargs):
            user_plan = user.get('plan', 'free')
            
            # Jerarquía de planes: free < growth < pro
            plan_hierarchy = {'free': 1, 'growth': 2, 'pro': 3}
            
            user_level = plan_hierarchy.get(user_plan, 1)
            required_level = plan_hierarchy.get(required_plan, 1)
            
            if user_level < required_level:
                return jsonify({
                    'error': 'plan_upgrade_required',
                    'current_plan': user_plan,
                    'required_plan': required_plan,
                    'message': f'Este feature requiere plan {required_plan}'
                }), 403
            
            return f(user, *args, **kwargs)
        return decorated_function
    return decorator

@app.route('/investors/<investor_id>/employees', methods=['GET'])
@require_auth
@require_plan('growth')
def get_fund_employees(user, investor_id):
    """
    Obtiene empleados de un fondo específico por investor_id
    Se usa cuando el usuario hace click en "Find Employees from fund"
    """
    try:
        print(f"🔍 Getting employees for investor: {investor_id}")
        
        # 1. Obtener información del investor
        investor_info = get_investor_by_id(investor_id)
        if not investor_info:
            return jsonify({
                "error": "Investor not found",
                "success": False
            }), 404
        
        company_name = investor_info.get('Company_Name')
        if not company_name:
            return jsonify({
                "error": "No company name found for this investor",
                "success": False
            }), 400
        
        print(f"🏢 Searching employees for company: {company_name}")
        
        # 2. Verificar créditos (costo fijo por búsqueda de empleados de un fondo)
        cost_per_search = 50  # Costo fijo por buscar empleados de un fondo
        user_credits = get_user_credits(user['id'])
        
        if user_credits < cost_per_search:
            return jsonify({
                "error": "insufficient_credits",
                "required": cost_per_search,
                "available": user_credits,
                "message": "Need 50 credits to find employees from fund"
            }), 402
        
        # 3. Buscar empleados de este fondo específico
        employees = find_employees_by_company_name(company_name)
        
        if not employees:
            return jsonify({
                "investor": {
                    "id": investor_id,
                    "company_name": company_name,
                    "description": investor_info.get('Company_Description', '')
                },
                "employees": [],
                "total_found": 0,
                "message": f"No employees found for {company_name}",
                "credits_charged": 0,
                "success": True
            })
        
        # 4. Cobrar créditos
        charge_success = charge_credits(user['id'], cost_per_search)
        if not charge_success:
            return jsonify({
                "error": "Could not charge credits"
            }), 500
        
        # 5. Formatear respuesta SIMPLIFICADA
        formatted_employees = format_fund_employees_simple(employees)
        
        result = {
            "investor": {
                "id": investor_id,
                "company_name": company_name,
                "description": investor_info.get('Company_Description', ''),
                "location": investor_info.get('Company_Location', ''),
                "linkedin": investor_info.get('Company_Linkedin', ''),
                "investing_stages": investor_info.get('Investing_Stage', ''),
                "investment_categories": investor_info.get('Investment_Categories', '')
            },
            "employees": formatted_employees,
            "total_found": len(formatted_employees),
            "credits_charged": cost_per_search,
            "credits_remaining": get_user_credits(user['id']),
            "success": True,
            "search_timestamp": datetime.now().isoformat()
        }
        
        print(f"✅ Found {len(formatted_employees)} employees for {company_name}")
        return jsonify(result)
        
    except Exception as e:
        print(f"❌ Error getting fund employees: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "error": "Could not get fund employees",
            "details": str(e),
            "success": False
        }), 500

def get_investor_by_id(investor_id):
    """Obtiene información del investor por ID"""
    try:
        query = """
        SELECT id, "Company_Name", "Company_Description", "Company_Location",
               "Investing_Stage", "Investment_Categories", "Company_Linkedin"
        FROM investors 
        WHERE id = %s
        """
        
        result = pd.read_sql(query, engine, params=[investor_id])
        
        if result.empty:
            return None
        
        return result.iloc[0].to_dict()
        
    except Exception as e:
        print(f"❌ Error getting investor: {e}")
        return None

def find_employees_by_company_name(company_name):
    """
    Encuentra TODOS los empleados de una empresa específica
    """
    try:
        query = """
        SELECT 
            id,
            "fullName" as full_name,
            "headline" as headline,
            "current_job_title" as job_title,
            "location" as location,
            "linkedinUrl" as linkedin_url,
            "email" as email,
            "profilePic" as profile_pic,
            "Company_Name" as company_name,
            "decision_score" as decision_score
        FROM employees 
        WHERE "Company_Name" = %s
        AND "fullName" IS NOT NULL 
        AND "fullName" != ''
        ORDER BY "decision_score" DESC NULLS LAST
        """
        
        result = pd.read_sql(query, engine, params=[company_name])
        
        if result.empty:
            return []
        
        # Convertir decision_score a numérico
        result['decision_score'] = pd.to_numeric(result['decision_score'], errors='coerce').fillna(0)
        
        return result.to_dict('records')
        
    except Exception as e:
        print(f"❌ Error finding employees: {e}")
        return []

def format_fund_employees_simple(employees):
    """
    Formatea empleados para el frontend - VERSIÓN SIMPLIFICADA
    """
    try:
        formatted = []
        
        for employee in employees:
            formatted_employee = {
                'id': str(employee['id']),
                'full_name': employee.get('full_name', ''),
                'headline': employee.get('headline', ''),
                'job_title': employee.get('job_title', ''),
                'location': employee.get('location', ''),
                'linkedin_url': employee.get('linkedin_url', ''),
                'email': employee.get('email', '') if employee.get('email') else None,
                'profile_pic': employee.get('profile_pic', ''),
                'company_name': employee.get('company_name', ''),
                'decision_score': float(employee.get('decision_score', 0))
            }
            
            formatted.append(formatted_employee)
        
        # Ordenar por decision_score (ya vienen ordenados de la query)
        return formatted
        
    except Exception as e:
        print(f"❌ Error formatting employees: {e}")
        return []

# Endpoint alternativo por Company_Name (si frontend prefiere esto)
@app.route('/companies/<company_name>/employees', methods=['GET'])
@require_auth 
@require_plan('growth')
def get_employees_by_company_name(user, company_name):
    """
    Alternativa: obtener empleados directamente por Company_Name
    """
    try:
        # Decodificar company_name si viene URL-encoded
        import urllib.parse
        company_name = urllib.parse.unquote(company_name)
        
        print(f"🔍 Getting employees for company: {company_name}")
        
        # Verificar créditos
        cost = 50
        user_credits = get_user_credits(user['id'])
        
        if user_credits < cost:
            return jsonify({
                "error": "insufficient_credits",
                "required": cost,
                "available": user_credits
            }), 402
        
        # Buscar empleados
        employees = find_employees_by_company_name(company_name)
        
        if not employees:
            return jsonify({
                "company_name": company_name,
                "employees": [],
                "total_found": 0,
                "message": f"No employees found for {company_name}",
                "credits_charged": 0,
                "success": True
            })
        
        # Cobrar créditos
        charge_credits(user['id'], cost)
        
        # Formatear y devolver
        formatted_employees = format_fund_employees_simple(employees)
        
        return jsonify({
            "company_name": company_name,
            "employees": formatted_employees,
            "total_found": len(formatted_employees),
            "credits_charged": cost,
            "credits_remaining": get_user_credits(user['id']),
            "success": True
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({
            "error": "Could not get employees",
            "success": False
        }), 500

def verify_investors_table():
    """Función para verificar que la tabla de inversores existe y tiene datos"""
    try:
        print("🔍 Verificando tabla de inversores...")
        
        with engine.connect() as conn:
            # Verificar si la tabla existe
            table_check = conn.execute(text("""
                SELECT table_name 
                FROM information_schema.tables 
                WHERE table_schema = 'public' 
                AND table_name = 'investors'
            """)).fetchone()
            
            if not table_check:
                print("❌ Tabla 'investors' no existe")
                return False
            
            print("✅ Tabla 'investors' existe")
            
            # Verificar contenido
            count_result = conn.execute(text("SELECT COUNT(*) FROM investors")).scalar()
            print(f"📊 Número de inversores en la tabla: {count_result}")
            
            if count_result == 0:
                print("⚠️ La tabla está vacía")
                return False
            
            return True
            
    except Exception as e:
        print(f"❌ Error verificando tabla: {e}")
        return False

# Verificar tabla al iniciar
if engine:
    verify_investors_table()
else:
    print("❌ Engine no disponible para verificación")

@app.route('/projects/<project_id>/save-investors', methods=['POST'])
@require_auth
def save_investors_to_project_endpoint(user, project_id):
    """
    API para guardar inversores encontrados en un proyecto
    El frontend envía la lista de inversores a guardar
    """
    try:
        data = request.get_json()
        
        if not data or 'investors' not in data:
            return jsonify({'error': 'Lista de inversores requerida'}), 400
        
        investors_list = data['investors']
        
        if not isinstance(investors_list, list):
            return jsonify({'error': 'investors debe ser una lista'}), 400
        
        # Verificar que el proyecto pertenece al usuario
        with engine.connect() as conn:
            project_check = conn.execute(
                text("SELECT 1 FROM projects WHERE id = :project_id AND user_id = :user_id"),
                {"project_id": project_id, "user_id": user['id']}
            ).fetchone()
            
            if not project_check:
                return jsonify({'error': 'Proyecto no encontrado'}), 404
        
        # Guardar inversores
        saved_count = save_investors_to_project_simple(user['id'], project_id, investors_list)
        
        return jsonify({
            'success': True,
            'message': f'{saved_count} inversores guardados exitosamente',
            'saved_count': saved_count,
            'project_id': project_id
        })
        
    except Exception as e:
        print(f"❌ Error saving investors: {e}")
        return jsonify({'error': 'No se pudieron guardar los inversores'}), 500
            
# ==============================================================================
#           MAIN
# ==============================================================================

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port, debug=False)
